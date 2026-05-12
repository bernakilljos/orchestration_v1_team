"""
build-arch-lecture-doc.py v2 — 19 챕터 × 8 섹션 표준.

각 챕터에 반드시 포함 (.claude/rules/teaching-doc.md):
  1. 📚 핵심 한 줄
  2. 📊 표 (비교·구조)
  3. 🌊 흐름도 / 단계
  4. 💪 강점
  5. ⚠️ 약점·주의
  6. ⭐ 강추 시점
  7. 🎯 우리 시스템 매핑 (orchestration_v1 의 어디·어떻게)
  8. 🧪 점검 1줄

이미지: 영어 원본 + 한글 다이어그램 (matplotlib 생성) 동반.
톤: 5살 청자, 친근 1인칭, 비유 풍부.
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent.parent
ARCH = ROOT / "docs" / "screens" / "arch"
ARCH_KOR = ROOT / "docs" / "screens" / "arch-kor"
OUT = ROOT / "docs" / "lecture-AI-Claude-초보자가이드.docx"


# ---------- 스타일 helpers ----------
def _set_kor(run, size=11, bold=False, color=None):
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rFonts.set(qn("w:ascii"), "맑은 고딕")
    rFonts.set(qn("w:hAnsi"), "맑은 고딕")


def P(doc, text="", size=11, bold=False, align="left", color=None, after=6):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(after)
    if text:
        _set_kor(p.add_run(text), size=size, bold=bold, color=color)
    return p


def H(doc, text, level=1, color=None):
    sizes = {1: 22, 2: 17, 3: 14, 4: 12}
    cdef = {1: (32, 56, 100), 2: (47, 84, 150), 3: (68, 114, 196), 4: (100, 100, 100)}
    return P(doc, text, size=sizes.get(level, 12), bold=True,
             color=color or cdef.get(level, (0, 0, 0)), after=10)


def B(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + 0.6 * level)
    p.paragraph_format.space_after = Pt(2)
    _set_kor(p.add_run(text), size=11)


def callout(doc, label, text, label_color=(192, 0, 0)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _set_kor(p.add_run(label + " "), size=11, bold=True, color=label_color)
    _set_kor(p.add_run(text), size=11)


class PageLayoutTracker:
    """페이지 콘텐츠 height 누적 — skill: auto-layout-fit."""
    PAGE_LIMITS = {"docx-landscape": 7.33, "docx-portrait": 9.5, "pptx-16:9": 6.7}
    H = {"h1": 0.55, "h2": 0.4, "callout": 0.5, "para_line": 0.18,
         "bullet_line": 0.2, "caption": 0.25, "table_row": 0.3, "safety": 0.3}

    def __init__(self, target="docx-landscape"):
        self.target = target
        self.used = 0.0

    def add(self, kind, count=1):
        self.used += self.H.get(kind, 0.2) * count

    def remaining(self):
        return self.PAGE_LIMITS[self.target] - self.used - self.H["safety"]

    def image_max_height(self):
        return max(2.0, self.remaining())

    def reset(self):
        self.used = 0.0


def IMG(doc, dir_path, name, width=5.8, height=None, caption=None, max_height=6.6):
    """이미지 임베드.

    max_height: 페이지 세로 한계 (landscape A4 = 약 7.0 inch, margin 빼면 6.6).
    width 비율 유지 시 height > max_height 면 height 기준으로 축소.
    """
    img = dir_path / name
    if not img.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)

    # PIL 로 실제 PNG 크기 읽어 비율 계산
    try:
        from PIL import Image as PILImage
        with PILImage.open(str(img)) as pim:
            iw, ih = pim.size
        # width=W inch 적용 시 height = W * (ih/iw) inch
        expected_h = width * (ih / iw)
        if expected_h > max_height:
            # height 기준으로 축소
            use_h = max_height
            use_w = max_height * (iw / ih)
            p.add_run().add_picture(str(img), width=Inches(use_w), height=Inches(use_h))
        else:
            p.add_run().add_picture(str(img), width=Inches(width))
    except Exception:
        # PIL 실패 시 fallback
        p.add_run().add_picture(str(img), width=Inches(width))

    if caption:
        P(doc, caption, size=9, align="center", color=(100, 100, 100))


def PB(doc):
    doc.add_page_break()


def HR(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def table(doc, header, rows, header_fill="4472C4"):
    """간단한 비교 표."""
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        _set_kor(hdr[i].paragraphs[0].add_run(h), size=10, bold=True, color=(255, 255, 255))
        tc = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), header_fill)
        tc.append(shd)
    for r, row in enumerate(rows):
        cells = t.rows[r + 1].cells
        for i, c in enumerate(row):
            cells[i].text = ""
            _set_kor(cells[i].paragraphs[0].add_run(str(c)), size=10)
    P(doc, "", after=4)


# ---------- 챕터 렌더링 ----------
CHAPTER_TO_KOR = {
    "1. AI 3종 세트": "01-gen-vs-agentic-vs-agent.png",
    "2. AI 에이전트의 8가지": "08-8-models.png",
    "3. 에이전트의 5가지": "02-5-cores.png",
    "4. 9가지 숨은 함정": "09-9-killers.png",
    "5. AI 스택 5층": "03-ai-stack-5layers.png",
    "6. 에이전트 개발킷": "04-dev-kit-5layers.png",
    "7. 제로비용 AI": "10-zero-cost.png",
    "8. AI 빌더 도구": "05-ai-builder-6cat.png",
    "9. RAG 입문": "11-rag-intro.png",
    "10. RAG 8가지": "12-rag-8.png",
    "11. API 프로토콜": "13-api-protocols.png",
    "12. MCP vs A2A": "06-mcp-vs-a2a.png",
    "13. Claude 마스터": "07-14-levels.png",
    "14. Claude Code 결정트리": "14-decision-tree.png",
    "15. Claude Code 완전 가이드": "15-complete-guide.png",
    "16. Claude Code 아키텍처 레퍼런스": "16-arch-reference.png",
    "17. Claude Code 프로젝트 구조": "17-project-structure.png",
    "18. .claude 폴더": "18-dk-folder.png",
    "19. CLAUDE.md 설계": "19-claude-md-design.png",
    "20. 8가지 프롬프트": "20-prompt-8.png",
}


def _kor_image_for(title):
    """챕터 제목에서 한글 PNG 매핑."""
    for prefix, png in CHAPTER_TO_KOR.items():
        if title.startswith(prefix):
            return png


# 사용자 전수조사 (2026-05-11) — 챕터별 이미지 max_height (inch)
# A4 landscape inside 7.33 - H1(0.55) - safety(0.3) = 6.48 한계 → max 6.4 clamp
# base = 6.0 → 사용자 ratio 그대로 적용 + 빈 페이지 방지
CHAPTER_MAX_HEIGHT = {
    "1. AI 3종 세트": 5.40,                    # -10%
    "2. AI 에이전트의 8가지": 6.00,            # 그대로
    "3. 에이전트의 5가지": 6.00,              # 그대로
    "4. 9가지 숨은 함정": 5.40,                # -10%
    "5. AI 스택 5층": 5.40,                   # -10%
    "6. 에이전트 개발킷": 5.40,                # -10%
    "7. 제로비용 AI": 5.40,                   # -10%
    "8. AI 빌더 도구": 5.40,                   # -10%
    "9. RAG 입문": 6.30,                       # +5%
    "10. RAG 8가지": 6.18,                     # +3%
    "11. API 프로토콜": 5.40,                  # -10%
    "12. MCP vs A2A": 6.06,                    # +1%
    "13. Claude 마스터": 6.00,                 # 재배치 (그대로)
    "14. Claude Code 결정트리": 6.00,          # 그대로 (좋아)
    "15. Claude Code 완전 가이드": 5.40,       # -10%
    "16. Claude Code 아키텍처 레퍼런스": 5.40, # -10%
    "17. Claude Code 프로젝트 구조": 5.40,     # -10%
    "18. .claude 폴더": 6.00,                  # 그대로 (좋아)
    "19. CLAUDE.md 설계": 5.40,                # -10%
    "20. 8가지 프롬프트": 5.40,                # -10%
}
DEFAULT_MAX_HEIGHT = 6.0


def _max_height_for(title):
    for prefix, mh in CHAPTER_MAX_HEIGHT.items():
        if title.startswith(prefix):
            return mh
    return DEFAULT_MAX_HEIGHT
    return None


def render_chapter(doc, ch, idx=0):
    """8 섹션 표준 챕터 렌더링 — PageLayoutTracker + 빈 페이지 방지.

    빈 페이지 방지 패턴: H1 paragraph 의 page_break_before 속성 사용.
    명시적 PB(doc) 는 Word 자동 분할과 충돌해 빈 페이지 생성 → 사용 금지.
    page_break_before 는 Word 가 H1 위치에서 직접 새 페이지 시작 → 빈 페이지 0.
    """
    tracker = PageLayoutTracker("docx-landscape")

    h1_para = H(doc, ch["title"], level=1)
    if idx > 0:
        h1_para.paragraph_format.page_break_before = True  # 두 번째 챕터부터 새 페이지
    # H1 + IMG 같은 페이지 보장 — keep_with_next + space_after 축소
    h1_para.paragraph_format.keep_with_next = True
    h1_para.paragraph_format.space_after = Pt(2)
    tracker.add("h1")

    # 사용자 전수조사 (2026-05-11) 챕터별 max_height 적용
    # H1 + IMG 한 페이지 + PB → callout/본문 별도 페이지 (잘림 방지)
    max_h = _max_height_for(ch["title"])
    kor_png = _kor_image_for(ch["title"])
    if kor_png:
        IMG(doc, ARCH_KOR, kor_png, width=10.5, max_height=max_h, caption=None)
    elif ch.get("image_kor"):
        IMG(doc, ARCH_KOR, ch["image_kor"], width=10.5, max_height=max_h, caption=None)

    # IMG 후 강제 PB — 본문 (callout + 표 + ...) 별도 페이지 → 잘림 방지
    PB(doc)
    callout(doc, "📚 핵심 한 줄", ch["핵심"]); tracker.reset(); tracker.add("callout")

    # 1.5 강사 한 마디 — 강사 톤 인트로 (친근·예시·학습자 가정)
    if ch.get("강사"):
        H(doc, "📢 강사 한 마디 (5살 청자에게 설명한다면)", level=3)
        for line in ch["강사"] if isinstance(ch["강사"], list) else [ch["강사"]]:
            P(doc, line, size=11, after=4)

    # 2. 표
    if ch.get("표"):
        H(doc, "📊 표 — 한눈에 비교", level=3)
        table(doc, ch["표"]["header"], ch["표"]["rows"])

    # 3. 흐름/단계
    if ch.get("흐름"):
        H(doc, "🌊 흐름·단계", level=3)
        for s in ch["흐름"]:
            B(doc, s)

    # 4. 강점
    if ch.get("강점"):
        H(doc, "💪 강점", level=3)
        for s in ch["강점"]:
            B(doc, s)

    # 5. 약점·주의
    if ch.get("약점"):
        H(doc, "⚠️ 약점·주의 (놓치기 쉬운 함정)", level=3)
        for s in ch["약점"]:
            B(doc, s)

    # 6. 강추 시점
    if ch.get("강추"):
        H(doc, "⭐ 강추 — 언제 써야 하나", level=3)
        if isinstance(ch["강추"], list):
            for s in ch["강추"]:
                B(doc, s)
        else:
            P(doc, ch["강추"])

    # 7. 우리 시스템 매핑 (가장 중요)
    if ch.get("우리시스템"):
        H(doc, "🎯 우리 시스템 매핑 — orchestration_v1 의 어디·어떻게", level=3)
        if isinstance(ch["우리시스템"], dict):
            table(doc, ["개념", "우리 시스템 위치 / 동작"], list(ch["우리시스템"].items()),
                  header_fill="C00000")
        else:
            for s in ch["우리시스템"]:
                B(doc, s)

    # 비유 (선택)
    if ch.get("비유"):
        H(doc, "🎯 비유 (5살에게 설명한다면)", level=3)
        P(doc, ch["비유"])

    # 8. 점검
    if ch.get("점검"):
        H(doc, "🧪 점검 — 한 줄 시험", level=3)
        q, a = ch["점검"]
        P(doc, "Q. " + q)
        P(doc, "A. " + a, color=(100, 100, 100))

    # 챕터 끝 PB 제거 — 다음 챕터의 시작 PB 가 처리 (빈 페이지 방지)


# ============================================================
# 챕터 데이터 (19개)
# ============================================================
CHAPTERS = [
    # ---- 1. Gen vs Agentic vs Agent ----
    {
        "title": "1. AI 3종 세트 — Generative vs Agentic vs AI Agent",
        "image_eng": "GenerativeAI-vs-AgenticAI-vs-AIAgents.jpg",
        "image_kor": "01-gen-vs-agentic-vs-agent.png",
        "핵심": "Generative = 글·그림을 한 번에 만들어 주는 AI. Agentic = 스스로 단계를 짜고 계획하는 AI. AI Agent = 외부 도구·API 까지 직접 호출하며 결과 평가하는 AI. 셋은 진화 단계이자, 한 시스템에 같이 들어갈 수도 있습니다.",
        "강사": ["여러분, AI 라고 하면 보통 ChatGPT 만 떠올리시죠? 근데 AI 가 사실 진화 단계가 있어요.", "Generative 는 '글만 쓰는 카피라이터'. 시키는 대로 한 번 답하고 끝.", "Agentic 는 '신입 매니저' — 목표 주면 단계를 짜요. 단 실행은 사람이.", "AI Agent 는 '경력 매니저' — 외부 API 까지 직접 호출하고 24/7 돌아갑니다.", "**우리 orchestration_v1 = Multi-Agent (가장 진화)** — Claude+Codex+Gemini+Haiku 가 회사처럼 협업."],
        "표": {
            "header": ["구분", "Generative", "Agentic", "AI Agent"],
            "rows": [
                ["하는 일", "한 번 답 만들고 끝", "스스로 단계 짜기", "API 호출까지 실행"],
                ["예시", "ChatGPT, DALL-E", "여행 계획 자동 설계", "비행기 예약 + 캘린더"],
                ["사람 비유", "카피라이터", "신입 매니저", "경력 매니저"],
                ["주도성", "낮음", "중간", "높음"],
                ["위험도", "낮음 (글뿐)", "중간", "높음 (실행)"],
            ],
        },
        "흐름": [
            "[Generative] 사용자 프롬프트 → 사전 학습된 거대 모델 (Transformer) → 토큰 시퀀스 예측 → 글·이미지·코드 출력. 한 번의 forward pass 로 끝, 자기 반성 없음.",
            "[Agentic] 사용자 목표 → LLM 이 작업 분해 (ReAct/CoT) → 각 단계마다 도구 선택 + 호출 → 중간 결과 검토 → 다음 단계 계획 → 최종 답. 단계는 LLM 이 직접 짬.",
            "[AI Agent] 목표 → 계획 + 메모리 로드 → 외부 API 실제 호출 (예약·이메일·DB) → 결과로 자기 평가 (성공/실패) → 보정 후 다음 행동 → 종료 조건까지 반복. 사람 개입 0.",
        ],
        "강점": [
            "[Generative] 빠르고 단순. 콘텐츠 (글·그림·코드) 즉시 생성. 학습 데이터 안에 있으면 매우 강함.",
            "[Generative] 비용 가장 낮음 (한 번 호출). 캐싱·배치 쉬움.",
            "[Agentic] LLM 이 계획을 명시적으로 펼쳐 보여줌 — 사람이 검토·수정 가능 (Human-in-the-Loop).",
            "[Agentic] 복잡한 다단계 작업도 분해 → 작업 단위로 검증 가능.",
            "[AI Agent] 사람 손 없이 끝까지 자동. 24/7 운영 가능.",
            "[AI Agent] 환경과 상호작용하며 학습·적응. 새 상황에도 대응.",
        ],
        "약점": [
            "[Generative] 시킨 것만 함. 모르면 자신만만하게 거짓말 (hallucination) — 9 함정 #9.",
            "[Generative] 단일 turn 한계. 긴 작업·다단계 추론 약함.",
            "[Agentic] 계획은 짜지만 실제 실행은 제한적 — 도구 결과 신뢰성 부담.",
            "[Agentic] 단계가 많아지면 비용 ↑ (호출 N 배), 지연 ↑.",
            "[AI Agent] 잘못된 실행 = 회복 어려움 (이메일 발송, DB 변경 등). 가드레일·롤백 필수.",
            "[AI Agent] 무한 루프 위험 (9 함정 #4). 예산·중지 조건 필수.",
        ],
        "강추": [
            "[Generative] 글·이미지·코드 한 번에 만들 때 — 블로그 글, 디자인 시안, 보일러플레이트 코드.",
            "[Agentic] 결과보다 '과정' 이 중요할 때 — 사용자에게 계획 보여주고 승인 받는 워크플로우.",
            "[Agentic] 도메인 전문가 + AI 협업 — AI 가 단계 짜고 사람이 검토·수정.",
            "[AI Agent] 반복 가능한 자동화 — 매일 같은 작업 (보고서·예약·모니터링).",
            "[AI Agent] hooks 가드레일 + observability 필수. 둘 다 갖춰야 운영 안전.",
        ],
        "우리시스템": {
            "어느 단계?": "AI Agent (가장 오른쪽). 더 정확히는 Multi-Agent System.",
            "Generative 능력": "Claude/Codex/Gemini 가 글·코드 생성",
            "Agentic 능력": "task-instruction.md 가 단계 정의 (단 자율 계획은 약함)",
            "AI Agent 능력": "MCP 도구 호출 + 자가 평가 (eval_quality) + 자가 복구 (watchdog)",
            "한 단계 더": "여러 워커 협력 = Multi-Agent (Claude→Codex→Gemini→Haiku 인수인계)",
        },
        "비유": "Generative = 카피라이터 한 명. Agentic = 신입 매니저. AI Agent = 경력 매니저. 우리 = AI 회사 한 채.",
        "점검": ("캘린더에 일정 자동 등록까지 해주는 AI 는?", "AI Agent. (Generative 는 글만, Agentic 는 계획만.)"),
    },
    # ---- 2. AI 8가지 모델 ----
    {
        "title": "2. AI 에이전트의 8가지 두뇌 — 모델 유형",
        "image_eng": "AI에이전트-8가지모델유형-jaiinfoway.jpg",
        "image_kor": None,
        "핵심": "AI 뇌는 한 종류가 아닙니다. 작업에 맞는 뇌를 골라 써야 비용·품질이 최적.",
        "강사": ["AI 뇌는 한 종류가 아니에요. 작업에 맞는 뇌를 골라써야 비용·품질이 최적이죠.", "비싼 모델 (Claude Opus) 만 쓰면 비용 폭증. 작은 모델 (Haiku) 로 80% 처리하고 큰 모델은 어려운 20% 만.", "**8개 모델 우리가 코드로 만드는 건 0개** — Claude·Codex·Gemini·Haiku 가 다 제공. 우리는 라우팅만 합니다.", "초보자 추천: GPT 계열 1개 (Claude or GPT-4) 로 시작 — 단순함 우선."],
        "표": {
            "header": ["뇌 유형", "예시", "★ 우리시스템 (현재)", "🔧 보완 (TODO)"],
            "rows": [
                ["GPT (Transformer)", "Claude", "Claude Opus 4.7 ✅", "—"],
                ["MoE (전문가 분배)", "Qwen 2", "route_dispatch.md (룰)", "★ 입력 자동 분류기 (1-2h)"],
                ["LRM (긴 추론)", "Gemini Flash", "Extended Thinking ✅", "Gemini quota 통합 (30m)"],
                ["VLM (이미지+글)", "GPT-4V", "Claude 멀티모달 ✅", "자동 OCR pipeline (1-2h)"],
                ["SLM (작은·빠름)", "Gemma", "Haiku 4.5 + cache 90%↓ ✅", "—"],
                ["LAM (행동)", "X-LAM", "MCP + Codex CLI ✅", "—"],
                ["HRM (계획→실행)", "Sapient", "auto-planner.md (skill)", "★ 자동 발동 강화 (30m)"],
                ["mHC (다층 협업)", "Deepseek", "킷 전체 구조 (수동 chain)", "자동 인수인계 chain (4-6h)"],
            ],
        },
        "흐름": [
            "사용자 요청 → 작업 유형 분류 (콘텐츠 생성 / 복잡 추론 / 도구 호출 / 이미지·멀티모달).",
            "분류 결과로 적합 모델 선택 — 비용·정확도·속도 우선순위 결정.",
            "단일 작업에 여러 모델 라우팅 가능 — 설계는 LRM, 구현은 GPT, 검증은 SLM 같은 분업.",
            "결과 검증 후 다음 작업 라우팅 — 자가 평가가 있으면 LAM/HRM 으로 보강.",
        ],
        "강점": [
            "비싼 모델만 안 쓰면 됨 — 작은 모델로 80% 처리, 큰 모델은 어려운 20% 만 호출 (비용 ↓ 60-80%).",
            "특정 작업 (이미지·행동·복잡 추론) 은 전용 모델 (VLM/LAM/LRM) 이 정확도·속도 ↑.",
            "라우팅 시스템 한 번 만들면 새 모델 추가 쉬움 — 인터페이스 표준화.",
            "각 모델 강점 합쳐 시너지 — GPT(글) + LAM(행동) + VLM(이미지) = 멀티모달 에이전트.",
            "Quota·rate limit 분산 가능 — 한 API 한계에 걸리면 다른 모델로 fallback.",
        ],
        "약점": [
            "모델 라우팅 로직 직접 만들어야 함 — 자동화 안 됨, 결정 트리·휴리스틱 필요.",
            "각 모델 API/토큰 비용·지연 다름 — 추적·관측 없으면 비용 폭증 (9 함정 #8).",
            "모델 간 출력 포맷 다름 — adapter·parsing 코드 추가 필요.",
            "라우팅 잘못하면 결과 품질 ↓ (예: 단순 작업에 LRM 쓰면 비싸고 느림).",
            "Vendor lock-in 위험 — 한 회사 모델만 쓰면 가격·정책 변경에 취약.",
            "버전 관리 복잡 — 각 모델 업데이트 영향 따로 검증해야.",
        ],
        "강추": [
            "처음엔 GPT 계열 1개 (Claude or GPT-4) 로 시작 — 단순함 우선.",
            "VLM 은 이미지·차트·UI 분석이 필요할 때만 추가 (비용 ↑).",
            "SLM 은 검증·필터링·분류 같은 단순 반복에 — Haiku, Gemma.",
            "LAM/HRM 은 도구 호출 많고 복잡한 워크플로우에서.",
            "사용자 요청 분류기 (route_dispatch) 만들어 자동 라우팅 — 우리 시스템 본보기.",
        ],
        "우리시스템": {
            "결론": "★ 8개 모델 우리가 구현 X — Claude·Codex·Gemini·Haiku 가 이미 제공. 우리는 '언제 어느 뇌 쓸지' 라우팅만.",
            "✅ 잘하는 3개": "GPT (Claude Opus 4.7), SLM (Haiku 4.5 + prompt cache 90%↓), LAM (MCP + Codex CLI)",
            "🔧 보완 필요 5개": "★ MoE 자동 분류기 (1-2h) · ★ HRM 자동 발동 (30m) · mHC 자동 인수인계 (4-6h) · LRM Gemini 통합 (30m) · VLM 자동 OCR (1-2h)",
            "🎯 가장 중요한 보완": "HRM 자동 발동 — 사용자가 매번 '5단계 해줘' 명시 X. Claude 가 자가 발동 (auto-planner skill description 매칭). 이번 세션 6시간 사용자 매번 지시받은 근본 원인.",
            "왜 보완?": "Generative→Agentic→Agent→Multi-Agent 진화 중 Agent 단계 약점. 자가 발동·자가 분배·자가 인수인계 부족 → 사용자 매번 지시.",
            "총 보완 시간": "1+2 우선 = 2시간 (즉시 큰 효과). 전부 = 8-10시간 (Phase 작업).",
        },
        "점검": ("8개 모델 중 우리가 코드로 구현하는 건 몇 개?", "0개. 모두 외부 모델 제공 — 우리는 라우팅·조율만."),
    },
    # ---- 3. 5 cores ----
    {
        "title": "3. 에이전트의 5가지 핵심 부품",
        "image_eng": "에이전틱AI-5가지핵심개념.jpg",
        "image_kor": "02-5-cores.png",
        "핵심": "안전한 에이전트는 5부품 모두 필요: 가드레일 / 오케스트레이션 / 도구(MCP) / 메모리 / 관측.",
        "강사": ["안전한 에이전트는 5부품 모두 필요해요. 회사로 비유하면:", "1) 입구 경비실 (Guardrails) — 잘못된 입력·민감정보 차단", "2) 부장님 (Orchestration) — 작업 분해·라우팅", "3) 도구함 (Tools/MCP) — GitHub·DB·API 안전 연결", "4) 서류 캐비넷 (Memory) — 단기/중기/장기 기억", "5) CCTV+회의록 (Observability) — 추적·비용 로그", "우리 시스템 = 5부품 모두 풍부. 24 hooks + 11 MCP + orca.db + recall-memory."],
        "표": {
            "header": ["부품", "역할", "회사 비유", "안 갖추면"],
            "rows": [
                ["1 Guardrails", "잘못된 입력·민감정보 차단", "입구 경비실", "사고·개인정보 유출"],
                ["2 Orchestration", "작업 분해·라우팅·상태관리", "부장님 작업 분배", "무한 루프·중복"],
                ["3 Tool & MCP", "외부 도구 안전 호출", "부서 도구함", "도구 못 씀"],
                ["4 Memory", "단기·중기·장기 기억", "서류 캐비넷", "맥락 망각"],
                ["5 Observability", "추적·비용·결정 로그", "CCTV+회의록", "왜 그랬는지 모름"],
            ],
        },
        "흐름": [
            "User Request → 1 Guardrails (입력 검증·PII 필터·rate limit) → 2 Orchestration (작업 분해·라우팅·상태) → 3 Tools/MCP (외부 도구 호출) → 4 Memory (단/중/장기 저장) → 5 Observability (트레이스·메트릭·로그) → Agent Response.",
            "5 레이어 간 context 전달 — 각 레이어가 다음 레이어에 정제된 데이터·결정 사유 전달.",
            "Observability 는 모든 레이어에서 trace 발생 — 사고 시 어느 레이어에서 문제인지 추적.",
        ],
        "강점": [
            "사고 났을 때 추적 가능 (5 Observability) — 어느 단계에서 망가졌는지 명확.",
            "외부 도구 호환 표준 (3 MCP) — 새 도구 추가 시 도구 자체만 만들면 됨, 시스템 안 건드림.",
            "메모리 분리로 맥락 안 깨짐 (4) — 단기 vs 장기 메모리가 서로 영향 X.",
            "각 레이어 독립 교체 가능 — Orchestration 만 LangGraph → CrewAI 로 갈아끼우기 OK.",
            "표준화로 팀 협업 쉬움 — 각 레이어 담당자 분리 가능.",
            "Guardrails 가 첫 방어선 — 시크릿 노출·PII 유출·악성 입력 차단.",
        ],
        "약점": [
            "5개 다 만들려면 초기 작업량 큼 — 작은 PoC 엔 과함.",
            "Observability 가 토큰·비용 ↑ — 모든 호출 로깅 시 스토리지 부담.",
            "레이어 간 인터페이스 표준화 안 되면 디버깅 어려움.",
            "Guardrails 너무 엄격하면 false positive — 정상 입력 차단.",
            "Memory 의 장기 (벡터 DB) 는 인덱스 구축 비용·시간 큼.",
            "Orchestration 의 상태 머신 복잡 → 버그 추적 어려움.",
        ],
        "강추": [
            "첫 프로젝트라도 Guardrails + Observability 두 개는 절대 빼지 마세요 — 사고 났을 때 추적 못 합니다.",
            "Memory 는 단기 (대화) + 중기 (세션) 두 층부터 — 장기 (벡터 DB) 는 필요할 때 추가.",
            "Orchestration 은 단순한 state machine 으로 시작 — 복잡한 게 좋은 게 아님.",
            "MCP 는 외부 도구 1개부터 — GitHub·Slack 같은 핵심부터 연결.",
        ],
        "우리시스템": {
            "1 Guardrails": "Hook (block_dangerous_bash·block-tricks·check-mojibake·protect-critical-files)",
            "2 Orchestration": "exec_orch plugin (route_dispatch·codex-auto·gemini-auto·watchdog)",
            "3 Tool & MCP": "MCP 서버 다수 (GitHub·Playwright·Figma·Slack·Notion 등)",
            "4 Memory": "orca.db (워커·태스크) + .claude/state/session-turns/ + ~/.claude/projects/<proj>/memory/",
            "5 Observability": "orca.db metrics + watchdog.log + external-watchdog 1분 간격 점검",
        },
        "비유": "에이전트 = 회사 한 채. 경비·부장·도구함·캐비넷·CCTV 다 갖춰야 사고 안 남.",
        "점검": ("우리 hooks 는 5부품 중 어디?", "1 Guardrails (입구 경비)."),
    },
    # ---- 4. 9 Silent Killers ----
    {
        "title": "4. 9가지 숨은 함정 — 에이전트가 망하는 패턴",
        "image_eng": "AI에이전트-9가지숨은킬러-판데이.jpg",
        "image_kor": None,
        "핵심": "성공한 에이전트는 다양하지만 망한 에이전트는 9가지 패턴 중 하나. 이걸 피하면 사고 80% ↓.",
        "강사": ["성공한 에이전트는 다양하지만 망한 에이전트는 9가지 패턴 중 하나. 이거 피하면 사고 80% ↓.", "가장 무서운 3개: **#4 Runaway Loop** (재시도 847번), **#8 Cost Blind** ($48,200 청구), **#9 No Failure Mode** (모르면 거짓말).", "우리 시스템 대응: watchdog backoff + orca.db budget + failure-mode.md 5중박기.", "9 중 8 ✅ 대응 — 본보기 수준입니다."],
        "표": {
            "header": ["#", "함정", "현상", "고치는 법"],
            "rows": [
                ["1", "Tool Bloat (도구 비대)", "도구 너무 많아 모델이 헷갈림", "도구 줄이고 날카롭게"],
                ["2", "Context Decay (맥락 부패)", "긴 대화에서 규칙 묻힘", "핵심 규칙 다시 박기"],
                ["3", "Retrieval Poisoning", "잘못 문서 들어옴 → 자신만만하게 틀림", "필터·랭킹·검증"],
                ["4", "Runaway Loop", "재시도 847번", "예산·중지조건·루프 가드"],
                ["5", "Schema Drift", "v1 → v2 필드 바뀐 거 못 봄", "스키마 버전·경계 검증"],
                ["6", "Eval Blindness", "10 예시로만 검증", "실 트래픽 슬라이스 평가"],
                ["7", "Non-Determinism", "같은 입력, 다른 답", "랜덤 제어·추적"],
                ["8", "Cost Blind", "어느날 $48,200 청구", "작업당 토큰·비용 추적"],
                ["9", "No Failure Mode", "모르면 거짓말", "거절·헤지 정책 박기"],
            ],
        },
        "흐름": [
            "프로덕션 가기 전 체크리스트 9개 출력 → 한 줄씩 점검 → 다 ✓ 면 배포 → 운영 중 매주 재점검.",
            "각 함정마다 자동 감시 메트릭 정의 — 임계 초과 시 알람·자동 차단.",
            "9 함정 중 #4 (Loop), #8 (Cost), #9 (Failure) 는 Hook 으로 강제 — AI 신뢰 X.",
        ],
        "강점": [
            "체크리스트 화 가능 — 한 번 만들면 모든 프로젝트에서 재사용.",
            "Hook 으로 #4 (루프) 와 #8 (비용) 강제 차단 가능 — AI 가 못 어김.",
            "운영 중 점검 가능 — 메트릭 임계로 자동 알람.",
            "팀 온보딩에 좋음 — 신입 개발자도 9 함정 알면 시야 ↑.",
            "9 함정 잡으면 PoC → 프로덕션 전환 시간 ↓.",
        ],
        "약점": [
            "#7 (비결정성) 는 LLM 본질이라 완전 제거 불가 — 추적·로깅이 최선.",
            "#6 (평가 맹점) 는 사람이 라벨 데이터 모아야 — 노동량 큼.",
            "#3 (Retrieval Poisoning) 은 RAG 사용 시만 — 단순 LLM 에는 무관.",
            "#9 (No Failure Mode) 의 거절 정책은 false positive 위험 — 정상 질문 거절.",
            "9 함정 다 잡으려면 초기 작업량 큼 — 작은 PoC 엔 과함.",
        ],
        "강추": [
            "#9 (No Failure) + #2 (Context Decay) 두 개만 잡아도 신뢰도 두 배.",
            "#4 (Loop) + #8 (Cost) 는 Hook 으로 강제 — 운영 안전망 필수.",
            "#7 (Non-Determinism) 은 추적 메트릭에 집중 — 완벽 제거 X.",
            "PoC 단계에선 #1·#2·#9 만, 프로덕션 가면 9 개 다 점검.",
        ],
        "우리시스템": {
            "#1 Tool Bloat": "개별 plugin (mcp_dev/mcp_data/mcp_collab 등) 으로 분리 — 필요한 것만 install",
            "#2 Context Decay": "CLAUDE.md + 5중 박기 + hook-00-init 매 세션 reminder 강제 출력",
            "#3 Retrieval Poisoning": "block-mojibake·check-korean-only 로 입력 검증",
            "#4 Runaway Loop": "watchdog backoff (10m→20m→40m→2h 지수) + quota_retry_count max 3 + kill_worker_tree",
            "#5 Schema Drift": "plugin.json schema v1.2 + validate-plugin-schema.py --strict",
            "#6 Eval Blindness": "eval_quality plugin + score-task (Haiku 자동 0-10 채점)",
            "#7 Non-Determinism": "orca.db metrics 에 모든 호출 기록 (model_id·tokens·cost·success)",
            "#8 Cost Blind": "orca.db budget 테이블 + daily_limit + breaker_tripped 자동 차단",
            "#9 No Failure Mode": "rules/failure-mode.md + 메모리 5중 박기 (\"모르면 거절\" 강제)",
        },
        "비유": "9 함정 = 회사 운영 9 사고. 우리는 9가지 다 hook·DB·rule 로 막아둠. 그래도 100% 는 아님 — 사람이 모니터링 필요.",
        "점검": ("\"$48,200 청구\" 사고는 9 함정 중?", "#8 Cost Blind. 우리는 orca.db budget 으로 차단."),
    },
    # ---- 5. AI 스택 5층 ----
    {
        "title": "5. AI 스택 5층 — 인프라부터 인터페이스까지",
        "image_eng": "AI스택-5단계-인프라부터인터페이스.jpg",
        "image_kor": "03-ai-stack-5layers.png",
        "핵심": "AI 시스템은 5층 빌딩. 아래 → 위: 인프라 → 데이터 → LLM → 오케스트레이션 → 인터페이스.",
        "강사": ["AI 시스템은 빌딩 같아요. 아래부터 5층 올리세요:", "1 Infra (실행) → 2 Data (RAG) → 3 LLM (추론) → 4 Orchestration (워크플로우) → 5 Interface (UI)", "**GPU 0 으로도 5층 다 작동** — LLM 만 외부 API (Claude) 쓰면 GPU 부담 0.", "우리 = 5/5 작동. Data 층 = ChromaDB 로컬 RAG 추가 (이번 세션)."],
        "표": {
            "header": ["층", "역할", "도구 예시"],
            "rows": [
                ["5 Interface", "사용자와 만남", "Streamlit·FastAPI·React·Vue·Auth0"],
                ["4 Orchestration", "여러 에이전트 흐름", "LangGraph·CrewAI·Claude Agent SDK"],
                ["3 LLM", "진짜 뇌", "Claude·GPT·Llama·OpenRouter"],
                ["2 Data", "벡터 DB·문서", "Chroma·Pinecone·Qdrant·Neo4j"],
                ["1 Infrastructure", "하드웨어·컨테이너", "Docker·K8s·AWS·GCP·RunPod"],
            ],
        },
        "흐름": [
            "사용자 → 5 Interface (UI) → 4 Orchestration (워크플로우) → 3 LLM (추론) → (필요시) 2 Data (RAG) → 1 Infra (실행) → 응답.",
            "각 층은 표준 인터페이스로 통신 — 다음 층 구현 알 필요 없음.",
            "확장 시 위에서부터 — 사용자 ↑ 면 Interface 부터, 정확도 ↑ 면 Data 강화, 비용 ↓ 면 LLM 라우팅.",
        ],
        "강점": [
            "계층 분리로 부분 교체 쉬움 — LLM 만 갈아끼우거나 Data 만 업그레이드 가능.",
            "각 층 독립 스케일링 — Interface 만 다중 인스턴스 띄울 수 있음.",
            "비용·성능 튜닝 명확 — 어느 층에 돈 쓸지 결정 가능.",
            "팀 분담 쉬움 — 인프라팀·데이터팀·ML팀·백엔드팀이 각자 층 담당.",
            "디버깅 명확 — 문제 어느 층인지 추적 가능.",
        ],
        "약점": [
            "스타트업은 5층 다 만들면 과중 — 클라우드 서비스로 1-2층 빌리는 게 보통.",
            "계층 간 데이터 복사 = 지연 ↑ (각 층 통과마다 직렬화·역직렬화).",
            "Vendor lock-in 위험 — 각 층 SaaS 의존 시 가격 인상에 취약.",
            "버전 호환성 — 한 층 업데이트가 다른 층에 영향 가능.",
            "Observability 가 5층 모두 커버해야 — 추적 시스템 복잡.",
        ],
        "강추": [
            "MVP 는 Streamlit (5) + Claude API (3) 만으로 — 30분 시작.",
            "트래픽 늘면 4 Orchestration 부터 추가 — LangGraph or CrewAI.",
            "데이터 검색 필요하면 2 Data — ChromaDB·Pinecone.",
            "프로덕션은 1 Infra 도 직접 — Docker + K8s.",
            "회사 정책상 클라우드 의존 X 면 Ollama (3) + ChromaDB (2) 로 로컬 풀스택.",
        ],
        "우리시스템": {
            "1 Infra": "사용자 PC 또는 VPS (Oracle Free Tier 권장 — exec_remote-setup)",
            "2 Data": "SQLite (orca.db) + 옵션 ChromaDB (exec_offline-vector)",
            "3 LLM": "Claude/Codex/Gemini/Haiku/로컬 LLM (라우팅)",
            "4 Orchestration": "exec_orch plugin (워커 + watchdog + 큐)",
            "5 Interface": "Claude Code CLI (terminal) + VS Code IDE",
        },
        "점검": ("스타트업 MVP 에 최소 어느 층?", "3 LLM + 5 Interface. 1·2·4 는 클라우드로 빌림."),
    },
    # ---- 6. 5 Dev Kit Layers ----
    {
        "title": "6. 에이전트 개발킷 5레이어 — CLAUDE.md + Skills + Hooks + Subagents + Plugins",
        "image_eng": "에이전트개발킷-5레이어구조-판데이.jpg",
        "image_kor": "04-dev-kit-5layers.png",
        "핵심": "Claude Code 의 5 부품. 이 5개를 잘 짜면 '확장·안전·팀배포' 다 됨.",
        "강사": ["Claude Code 의 5 레이어 = CLAUDE.md + Skills + Hooks + Subagents + Plugins.", "각 레이어 책임 분명 — 어디서 결정됐는지 명확.", "**Hooks 가 결정론적 강제** — AI 신뢰 X. 위험 명령 차단은 무조건 Hook.", "우리 = **34 hooks + 145 skills + 44 subagents + 26 plugins** — 본보기 수준."],
        "표": {
            "header": ["Layer", "이름", "역할", "비유"],
            "rows": [
                ["1", "CLAUDE.md", "메모리·규칙", "사규집"],
                ["2", "Skills", "지식·워크플로우", "부서 매뉴얼"],
                ["3", "Hooks", "가드레일", "보안 검색대"],
                ["4", "Subagents", "위임·격리", "특수 부서"],
                ["5", "Plugins", "묶음 배포", "부서 패키지"],
            ],
        },
        "흐름": [
            "1 CLAUDE.md (큰 규칙 정의) → 2 Skills (작업별 전문지식) → 3 Hooks (안전 강제) → 4 Subagents (큰 작업 위임) → 5 Plugins (1-4 묶음 + 팀 배포).",
            "레이어 마다 독립 진화 — 새 Skill 추가해도 CLAUDE.md 안 건드림.",
            "Plugin 은 marketplace 또는 git 으로 공유 — 한 번 만들면 팀 전체 자동 적용.",
        ],
        "강점": [
            "각 레이어 책임 분명 — 디버깅 쉽고 변경 영향 적음.",
            "Plugin 으로 팀·다른 프로젝트에 통째 배포 (우리 install/setup 패턴 본보기).",
            "Hooks 는 결정론적 강제 — AI 신뢰성 무관 (예: 위험 명령 차단).",
            "Subagents 격리로 메인 컨텍스트 보호 — 큰 탐색·테스트 위임.",
            "CLAUDE.md 의 3 scope (Global/Project/Folder) 로 규칙 충돌 자동 해결.",
            "Skills 자동 호출 — description 매칭만 잘 쓰면 사용자 명시 호출 X.",
        ],
        "약점": [
            "레이어 많으면 시작 진입장벽 ↑ — 처음엔 CLAUDE.md + Skills 두 개로 시작 추천.",
            "Subagent 는 컨텍스트 격리지만 호출 비용 ↑ — 자체 모델·도구 인스턴스.",
            "Hooks 는 디버깅 어려움 — 자동 발동이라 어느 hook 이 차단했는지 추적 필요.",
            "Skills 가 description 매칭에 실패하면 호출 X — false negative.",
            "Plugin 의존성 관리 — 한 plugin 이 다른 plugin 의존하면 순서 중요.",
        ],
        "강추": [
            "혼자 쓸 거면 Layer 1·2 만으로도 충분 — CLAUDE.md + 기본 skill 몇 개.",
            "팀 작업 = 무조건 5 Plugin 화 — git 으로 공유, install 로 자동 배포.",
            "위험 차단 (rm -rf, 시크릿 노출) = 반드시 3 Hook — AI 신뢰 X.",
            "큰 코드베이스 탐색 = 4 Subagent (explorer) — 메인 컨텍스트 보호.",
            "재사용 워크플로우 (PR 리뷰·테스트) = 2 Skill 화.",
        ],
        "우리시스템": {
            "1 CLAUDE.md": "프로젝트 CLAUDE.md (~169줄) + ~/.claude/CLAUDE.md (글로벌, install 자동배포)",
            "2 Skills": ".claude/skills/ 77개 (예: arch-mindmap, design_ppt, route_dispatch)",
            "3 Hooks": "PreToolUse + PostToolUse + SessionStart + Stop + SessionEnd (24개 등록)",
            "4 Subagents": ".claude/agents/ 11개 (team-lead, reviewer, judge, explorer, test-runner)",
            "5 Plugins": "plugins/ 25개 (exec_orch, eval_quality, mcp_*, design_*, etc)",
        },
        "비유": "회사 시스템. 사규(1) → 매뉴얼(2) → 검색대(3) → 특수부서(4) → 부서패키지(5). 사규부터 정해야 모두 같은 방향.",
        "점검": ("위험 명령 차단은 5레이어 중?", "3 Hooks. AI 가 아니라 결정론적 강제."),
    },
    # ---- 7. Zero-cost 2026 ----
    {
        "title": "7. 제로비용 AI 아키텍처 2026 — 0원으로 만드는 AI",
        "image_eng": "제로비용-AI아키텍처-2026-판데이.jpg",
        "image_kor": None,
        "핵심": "회사 카드 없어도 AI 서비스 띄울 수 있음. 무료 티어 + 로컬 LLM 조합.",
        "강사": ["회사 카드 없어도 AI 띄울 수 있어요. 무료 도구만으로:", "Frontend (Next.js/Streamlit) · Orchestrator (LangGraph) · LLM (Ollama+Gemma 로컬) · RAG (ChromaDB) · Tool (MCP) · Deploy (Cloudflare)", "**우리 = 거의 제로비용 본보기** — Claude Code CLI + exec_orch + SQLite + 선택 Ollama. 유료는 Claude API ($20/월) 만.", "초보자: 1+2 만으로 80% 작업. 작은 모델 + 작은 폴더로 시작하세요."],
        "표": {
            "header": ["층", "무료 도구", "유료 대체"],
            "rows": [
                ["Frontend", "Streamlit / Vercel free", "Vercel Pro"],
                ["Orchestrator", "LangGraph / CrewAI", "Claude Agent SDK"],
                ["LLM (로컬)", "Ollama + Gemma/Llama", "Claude / GPT API"],
                ["RAG", "LlamaIndex + ChromaDB", "Pinecone"],
                ["Tool Use", "MCP (open)", "유료 SaaS API"],
                ["DB", "SQLite / Supabase free", "Postgres 유료"],
                ["Observability", "Phoenix self-host", "LangSmith / Datadog"],
                ["Deploy", "Docker / HF Spaces", "AWS / GCP 유료"],
            ],
        },
        "흐름": [
            "사용자 → Vercel Frontend (free) → LangGraph Orchestrator → Ollama LLM (로컬, 무료) → 필요시 RAG (LlamaIndex + ChromaDB) → SQLite Data → 응답.",
            "Phoenix self-host 가 모든 단계 관측 — 어디서 막혔는지 추적.",
            "Docker / Cloudflare Workers / HF Spaces 로 무료 배포.",
        ],
        "강점": [
            "금전 비용 0 — GPU 만 있으면 (개인 PC 활용).",
            "오픈소스라 커스터마이즈 자유 — 코드 수정 가능.",
            "데이터·모델 자기 통제 — privacy·security 강함.",
            "Vendor lock-in 없음 — 언제든 갈아끼움.",
            "학습용 최고 — 각 층 직접 만들어 보며 이해.",
        ],
        "약점": [
            "GPU·CPU 본인 자원 — 트래픽 늘면 한계 (사용자 100+ 면 부담).",
            "운영 복잡도 ↑ — 직접 모니터링·업데이트·백업.",
            "지원·SLA 없음 — 사고 나면 본인이 처리.",
            "최신 모델 못 씀 — Claude 4.x·GPT-4 같은 SOTA 는 API 만.",
            "초기 셋업 시간 ↑ — Ollama·ChromaDB·LangGraph 다 설치·튜닝.",
        ],
        "강추": [
            "PoC·개인 프로젝트·학습용 — 최적.",
            "사내 폐쇄망 (인터넷 X) 환경 — 외부 API 못 쓸 때 유일.",
            "트래픽 늘면 LLM 부분만 유료 API 로 갈아끼움 — 부드러운 진화.",
            "RAG 만 로컬 (사내 문서 보안) + LLM 은 API 하이브리드 추천.",
            "오프라인 데모·강의 — Ollama 면 인터넷 없이도.",
        ],
        "우리시스템": {
            "Frontend": "Claude Code CLI (terminal) — 무료",
            "Orchestrator": "exec_orch (자체 구현) — 무료",
            "LLM (로컬)": "exec_offline-setup 으로 Ollama 설치 옵션",
            "RAG": "spec-only — 향후 ai_rag plugin",
            "DB": "SQLite (orca.db) — 무료",
            "Observability": "watchdog.log + external-watchdog — 무료",
        },
        "점검": ("로컬 LLM 무료 옵션 한 줄?", "Ollama + Gemma 또는 Llama. GPU 만 있으면 0원."),
    },
    # ---- 8. AI 빌더 매트릭스 ----
    {
        "title": "8. AI 빌더 도구 매트릭스 — 30가지 도구를 6 카테고리로",
        "image_eng": "AI빌더도구-6카테고리매트릭스.jpg",
        "image_kor": "05-ai-builder-6cat.png",
        "핵심": "AI 도구 백화점. 본인 필요 카테고리만 골라 둘러보기.",
        "강사": ["AI 도구 백화점. 6 카테고리 × 5 도구 = 30 개 중 본인 필요한 거만 골라쓰세요.", "초보 = 1+2 (Claude + Claude Code) 만으로 80% 작업.", "MVP 만들기 = 3 (Lovable·v0) 추가.", "회사 배포 = 4·5 (HuggingFace·LangChain).", "콘텐츠 = 6 (Mirra·Midjourney·Runway).", "우리 = 5/6 활용 + 1개 의도적 미사용 (앱·프로토타입은 우리 작업 외)."],
        "표": {
            "header": ["#", "카테고리", "베스트 도구"],
            "rows": [
                ["1", "모델·검색", "ChatGPT / Claude / Gemini / Perplexity / Grok"],
                ["2", "코딩·에이전트", "Cursor / Claude Code / Windsurf / Copilot / Replit"],
                ["3", "앱·프로토타입", "Lovable / Bolt / v0 / Framer AI / Vercel SDK"],
                ["4", "데이터·인프라", "HuggingFace / Replicate / Modal / RunPod / Pinecone"],
                ["5", "워크플로우", "LangChain / LlamaIndex / n8n / Make / Browserbase"],
                ["6", "미디어·콘텐츠", "★Mirra / Midjourney / Runway / ElevenLabs / ComfyUI"],
            ],
        },
        "흐름": [
            "요구사항 정의 (콘텐츠? 코딩? 앱? 데이터?) → 적합 카테고리 선택 → 카테고리 안 1-2 도구 무료 시도 → 가장 맞는 거 정착 → 부족하면 옆 카테고리 추가.",
            "초기엔 1+2 (모델+코딩) 만으로 80% — 다른 카테고리는 필요할 때 추가.",
            "도구 간 통합 (예: Cursor + Claude + n8n) 으로 워크플로우 확장.",
        ],
        "강점": [
            "한 장에 30 도구 매핑 — 의사결정 빠름.",
            "카테고리 별로 무료/유료 옵션 다양 — 예산 맞춤.",
            "최신 트렌드 반영 — 2026 기준 가장 핫한 도구 (Cursor·Mirra·Lovable).",
            "단계별 진화 가능 — 초보→중급→전문 단계 별 추천 명확.",
            "한 도구가 fail 해도 같은 카테고리 다른 도구로 즉시 교체.",
        ],
        "약점": [
            "도구 너무 많으면 결정 마비 — 처음엔 카테고리 1+2 만.",
            "최신 도구 (Mirra·v0 같은) 는 안정성 검증 필요 — 운영 위험.",
            "카테고리 간 통합 표준 없음 — 자체 glue code 필요.",
            "30 도구 다 라이센스·계정 관리 부담 — 비용·보안.",
            "도구 빠른 변화 — 6개월마다 매트릭스 재검토 필요.",
        ],
        "강추": [
            "초보 = 1(Claude) + 2(Claude Code) 로 80% 작업.",
            "MVP 만들기 = 1 + 2 + 3 (Lovable or v0).",
            "회사 배포 = 4 + 5 도 같이.",
            "콘텐츠 크리에이터 = 6 만.",
        ],
        "우리시스템": {
            "1 모델": "Claude (Opus/Sonnet/Haiku) — 메인 LLM",
            "2 코딩": "Claude Code (이 도구!) + codex-auto + gemini-auto",
            "3 앱": "사용 X (필요시 Streamlit 추천)",
            "4 데이터": "SQLite (orca.db) — 작은 규모",
            "5 워크플로우": "exec_orch (자체 구현) — LangChain 대체",
            "6 미디어": "design_ppt + design_word + design_excel (자체)",
        },
        "점검": ("코드 모르고 앱 만들고 싶다 — 어느 카테고리?", "3 앱·프로토타입. Lovable·Bolt·v0 으로 프롬프트 → 앱."),
    },
    # ---- 9. RAG 입문 ----
    {
        "title": "9. RAG 입문 — '읽고 답하는 AI'",
        "image_eng": "RAG-고전vs그래프vs에이전틱-판데이.jpg",
        "image_kor": None,
        "핵심": "RAG = 검색(Retrieval) + 생성(Generation). LLM 이 모르는 내용을 외부 문서에서 찾아 답하게 함.",
        "강사": ["RAG = 검색(Retrieval) + 생성(Generation). LLM 이 모르는 내용을 외부 문서에서 찾아 답하게 합니다.", "장점: 출처 표시 가능 (감사·검증 강함) + Hallucination ↓.", "**우리 = ChromaDB + 한글 임베딩 (paraphrase-multilingual)** + 36 docs indexed.", "Hook 자동 통합 — 사용자 메시지 시 관련 memory 자동 recall."],
        "표": {
            "header": ["방식", "흐름", "정확도", "복잡도"],
            "rows": [
                ["Classic", "질문 → 임베딩 → 벡터DB Top-K → LLM", "보통", "낮음"],
                ["Graph", "엔티티 추출 → 지식그래프 → LLM", "관계 강함", "중간"],
                ["Agentic", "추론 에이전트 → 다중 소스 → 자가 검증", "최고", "높음"],
            ],
        },
        "흐름": [
            "사내 문서·웹 → 청크 분할 → 임베딩 모델 → 벡터 DB 인덱싱 (사전 작업).",
            "사용자 질문 → 임베딩 (수치 변환) → 벡터 DB 검색 (코사인 유사도 Top-K) → 관련 문서 추출.",
            "추출 문서 + 원 질문 → LLM 프롬프트 → 답변 생성 + 출처 인용.",
            "Agentic RAG 는 추가로 자가 평가 (검색 결과 충분한지) → 부족하면 재검색.",
        ],
        "강점": [
            "LLM 이 학습 안 한 최신·사내 정보도 답할 수 있음 — 회사 위키 검색 같은 거.",
            "출처 표시 가능 — 어느 문서에서 왔는지 명확 (감사·검증 강함).",
            "LLM 재학습 (fine-tuning) 보다 싸고 빠름 — 문서 추가만 하면 즉시 반영.",
            "Hallucination 줄임 — 실제 문서 근거로 답.",
            "도메인 특화 쉬움 — 법률·의료·사내 매뉴얼 같은 닫힌 영역에 강함.",
        ],
        "약점": [
            "검색 품질이 답 품질 결정 — 잘못 검색 = 자신만만 틀림 (9 함정 #3).",
            "벡터DB 인덱스 구축 시간·비용 — 100만 문서면 인덱싱만 몇 시간.",
            "임베딩 모델 선택 중요 — 도메인 안 맞으면 검색 실패.",
            "청크 크기 튜닝 어려움 — 너무 크면 검색 부정확, 작으면 맥락 부족.",
            "최신 정보 반영 지연 — 새 문서 추가 → 재인덱싱 → 검색 가능.",
        ],
        "강추": [
            "사내 위키·매뉴얼·법규 문서 같은 닫힌 도메인 — RAG 최적.",
            "고객지원 챗봇 — FAQ + 매뉴얼 RAG.",
            "법률·의료·금융 — 출처 표시 필수 영역.",
            "처음엔 Naive RAG (단순 벡터) 로 — 정확도 부족하면 Corrective 추가.",
            "LLM 재학습 (fine-tuning) 비싸고 느리면 RAG 가 답.",
        ],
        "우리시스템": {
            "현재 상태": "✅ RAG 구현 완료 (ai_rag plugin stable v1.0)",
            "도구": "ChromaDB + sentence-transformers (paraphrase-multilingual-MiniLM-L12-v2 — 한·영 모두)",
            "indexed docs": "36개 (CLAUDE.md · .claude/rules/ · memory/feedback · plugins/skills)",
            "구현 방식": "Naive (벡터) + Hybrid (kw+rag) — hook 자동 통합",
            "자동 재빌드": "feedback/rule/skill md 변경 시 PostToolUse hook → ChromaDB 자동 rebuild",
            "8 아키텍처 중": "Naive ✅, Hybrid ✅, 나머지 6 (Corrective/HyDE/Graph/Adaptive/Agentic/Multimodal) plugin spec",
        },
        "점검": ("LLM 이 모르는 사내 정책을 묻는다면?", "RAG 필요. 사내 문서 인덱싱 후 검색 + LLM."),
    },
    # ---- 10. RAG 8 architectures ----
    {
        "title": "10. RAG 8가지 아키텍처",
        "image_eng": "RAG-8가지아키텍처-dailydoseofds.jpg",
        "image_kor": None,
        "핵심": "내 데이터 특성·정확도 요구에 따라 RAG 도 8가지 중 골라야.",
        "강사": ["데이터 특성·정확도 요구에 따라 RAG 도 8가지 중 골라야:", "Naive (PoC) → Corrective (정확도) → HyDE (모호) → Graph (관계) → Adaptive (다양) → Agentic (최고).", "**우리 = 8/8 모두 구현** (rag-naive·corrective·hyde·adaptive·agentic·graph·multimodal·hybrid). 학계 표준 다 작동."],
        "표": {
            "header": ["RAG 종류", "특징", "강추 상황"],
            "rows": [
                ["Naive", "기본형, 단순·빠름", "PoC, 단순 사실 조회"],
                ["Multimodal", "이미지+텍스트", "도식·표가 많은 문서"],
                ["HyDE", "가상 답 만들고 검색", "모호한 질문"],
                ["Corrective", "결과 채점·웹 fallback", "정확도 매우 중요"],
                ["Graph", "지식 그래프", "관계·다중 출처"],
                ["Hybrid", "벡터+그래프 동시", "최고 정확도, 복잡"],
                ["Adaptive", "질문 분석 후 분기", "다양한 질문 유형"],
                ["Agentic", "ReAct + 멀티 에이전트", "복잡 추론·도구 같이"],
            ],
        },
        "흐름": [
            "Naive RAG 시작 (가장 단순) → 정확도 부족하면 Corrective (결과 채점) 또는 HyDE (가상답으로 검색) 추가.",
            "다중 출처·관계 검색 필요 → Graph RAG (지식 그래프) 도입.",
            "최고 정확도 필요 → Hybrid (Vector+Graph) 또는 Agentic (자율 검증 루프).",
            "Multimodal (이미지+텍스트) 필요 → 별도 분기.",
        ],
        "강점": [
            "선택지 풍부 — 상황별 최적화 가능.",
            "단계별 진화 가능 — Naive → Corrective → Graph → Agentic 점진 강화.",
            "각 방식이 다른 한계 보완 — 조합 가능.",
            "오픈소스 풍부 (LlamaIndex·LangChain 등) — 구현 부담 ↓.",
            "8 방식 다 표준 패턴 — 학계·업계 검증 완료.",
        ],
        "약점": [
            "선택 어려움 — 처음엔 Naive 면 충분, 과한 설계 함정.",
            "8 방식마다 구축 비용·복잡도 다름 — Hybrid·Agentic 은 매우 큼.",
            "Agentic RAG 는 느림·비용 ↑ — 자가 검증 루프 토큰 소비.",
            "Multimodal 은 임베딩 모델 부담 — 이미지·텍스트 통합 모델 필요.",
            "Adaptive RAG 의 질문 분류기 만들기 자체 노동.",
        ],
        "강추": [
            "PoC = Naive — 1주일 안에 검증.",
            "정확도 ↑ 필요 = Corrective + HyDE — 답 품질 두 배.",
            "관계·다중 출처 검색 = Graph RAG.",
            "최고 품질 + 비용 OK = Hybrid (Vector+Graph) 또는 Agentic.",
            "복잡 추론 (Multi-hop) = Agentic — Self-Eval 포함.",
        ],
        "우리시스템": {
            "결론": "★ 8/8 모두 구현 완료 (2026-05-12 세션)",
            "Naive": ".claude/scripts/rag-recall.py — ChromaDB 기본 벡터 검색",
            "Multimodal": "rag-multimodal.py — 이미지 캡션 + 텍스트 통합 검색",
            "HyDE": "rag-hyde.py — 가상 답 4개 생성 후 검색",
            "Corrective": "rag-corrective.py — distance > 0.7 약함 → query 변형 재검색",
            "Graph": "rag-graph.py — entity 추출 + co-occurrence graph (orca.db graph_edges 1059 edges)",
            "Hybrid": "user-prompt-auto-planner.sh hook 안 kw+rag 동시 (keyword + 의미)",
            "Adaptive": "rag-adaptive.py — factual/complex/vague/multi_hop 자동 분류 → 적합 RAG 분기",
            "Agentic": "rag-agentic.py — Corrective + 자가 평가 (confidence) + HyDE fallback",
        },
        "점검": ("회사 정책 + 외부 법규 같이 검색하려면?", "Hybrid 또는 Agentic. 우리 시스템 = 둘 다 구현됨."),
    },
    # ---- 11. API Protocols ----
    {
        "title": "11. API 프로토콜 한눈에 보기",
        "image_eng": "API프로토콜-한눈에.jpg",
        "image_kor": None,
        "핵심": "외부 시스템 연결 약속 11가지. AI 작업엔 거의 REST + Webhooks 만 알면 됨.",
        "강사": ["프로토콜 11가지 중 AI 작업엔 **REST + Webhooks 만 알면 80%** 커버.", "REST = HTTP+JSON 표준. 거의 모든 API.", "Webhooks = 이벤트 콜백. GitHub·Slack 알림.", "토큰 스트리밍 = SSE (Claude API 자동).", "우리 = REST 11+ MCP + Webhooks (GitHub/Slack) + EDA (auto-dispatch)."],
        "표": {
            "header": ["프로토콜", "쓰는 곳", "한 줄"],
            "rows": [
                ["REST", "거의 모든 API", "HTTP+JSON, 표준"],
                ["Webhooks", "이벤트 알림", "서버 → 클라이언트 콜백"],
                ["GraphQL", "복잡 UI", "원하는 필드만 요청"],
                ["WebSocket", "채팅·게임", "양방향 실시간"],
                ["SSE", "실시간 푸시", "서버 → 클라이언트 일방향"],
                ["gRPC", "고성능 내부 통신", "Protocol Buffers"],
                ["SOAP", "은행·정부", "XML 기반 전통"],
                ["AMQP/MQTT", "메시지 큐·IoT", "비동기 메시징"],
                ["EDA", "마이크로서비스", "이벤트 발행/구독"],
                ["EDI", "기업 간 문서", "발주서·인보이스"],
            ],
        },
        "흐름": [
            "요구사항 분류 (요청-응답·이벤트·스트림·메시지큐) → 적합 프로토콜 선택.",
            "AI 도구 호출 = 거의 REST. 외부 알람 받기 = Webhooks. 토큰 스트리밍 = SSE.",
            "마이크로서비스 = EDA + AMQP. IoT = MQTT. 내부 고성능 = gRPC.",
        ],
        "강점": [
            "REST 한 가지만 마스터해도 80% 커버 — 대부분 AI 작업 충분.",
            "표준화 명확 — 도구 간 호환성 ↑.",
            "각 프로토콜 별 SDK 풍부 — 구현 부담 ↓.",
            "필요할 때만 다른 프로토콜 추가 — 점진 확장.",
            "Webhooks + REST 조합으로 양방향 통신 가능.",
        ],
        "약점": [
            "프로토콜 많으면 결정 피로 — 단순한 것부터.",
            "REST 외엔 학습 곡선 ↑ — gRPC·SOAP 는 진입장벽.",
            "WebSocket·SSE 는 연결 관리 복잡 — heartbeat·재연결 처리.",
            "AMQP·MQTT 는 메시지큐 인프라 필요 — RabbitMQ·Mosquitto 운영 부담.",
            "프로토콜 혼용 시 디버깅 어려움 — 추적 시스템 통합 필요.",
        ],
        "강추": [
            "처음엔 REST + JSON — 80% 완료.",
            "실시간 알림 필요 = Webhooks (서버→클라이언트).",
            "토큰 스트리밍 (LLM 답변) = SSE — 가장 단순한 단방향 push.",
            "양방향 실시간 (채팅) = WebSocket.",
            "마이크로서비스 = EDA + AMQP — Kafka·RabbitMQ.",
        ],
        "우리시스템": {
            "REST": "MCP 서버들이 내부적으로 REST 사용",
            "Webhooks": "GitHub MCP·Slack MCP 가 webhook 콜백 받음",
            "기타": "AI 작업 범위에서 거의 REST + Webhooks 로 충분",
        },
        "점검": ("실시간 알림이 필요하면?", "Webhooks (서버에서 푸시) 또는 SSE."),
    },
    # ---- 12. MCP vs A2A ----
    {
        "title": "12. MCP vs A2A — 두 프로토콜의 차이",
        "image_eng": "MCP-vs-A2A-프로토콜비교-datasciencedojo.jpg",
        "image_kor": "06-mcp-vs-a2a.png",
        "핵심": "MCP = LLM이 도구를 직접 부름. A2A = 매니저가 부서장(에이전트)에게 위임.",
        "강사": ["MCP = LLM 한 명이 도구 호출. A2A = 여러 Agent 가 자율 협력.", "보통 = MCP (단순·디버깅 쉬움). 멀티 도메인 = A2A.", "**우리 = MCP 11 서버 + A2A-lite chain** (Claude 통제 + auto-dispatch). 둘 다 본보기."],
        "표": {
            "header": ["구분", "MCP", "A2A"],
            "rows": [
                ["주체", "LLM 한 명", "여러 에이전트"],
                ["호출 대상", "외부 도구 (API)", "다른 에이전트"],
                ["제어권", "LLM이 끝까지 통제", "각 에이전트가 자율"],
                ["비유", "비서 + 도구함", "매니저 + 부서장"],
                ["복잡도", "낮음", "높음"],
                ["언제", "도구 표준화 필요", "도메인별 자율 협력 필요"],
            ],
        },
        "흐름": [
            "[MCP] User → LLM → MCP Client (도구 선택) → MCP Server (도구 wrapper) → 외부 API → 결과 → LLM → 사용자 응답. LLM 이 끝까지 통제.",
            "[A2A] User → Orchestrator Agent (작업 분해) → 도메인 Agent (위임) → Agent 가 자체 도구 호출 → 결과 → Orchestrator → 통합 응답. 각 Agent 자율.",
            "혼합: MCP 로 도구 표준화 + A2A 로 Agent 협력 — 우리 시스템 패턴.",
        ],
        "강점": [
            "[MCP] 단순·검증 쉬움 — LLM 한 명 통제, 흐름 추적 명확.",
            "[MCP] 도구 표준화로 새 API 추가 쉬움 — Wrapper 만 만들면 됨.",
            "[A2A] 확장 좋음 — 새 도메인 Agent 추가가 시스템 영향 적음.",
            "[A2A] 자율성 ↑ — 각 Agent 가 자기 영역 책임.",
            "[A2A] 멀티 도메인 협력 자연 — 항공·호텔·캘린더 Agent 따로 자율 결정.",
        ],
        "약점": [
            "[MCP] LLM 컨텍스트 폭주 위험 — 도구 결과 전부 LLM 에 들어감.",
            "[MCP] LLM 의존도 ↑ — 모델 능력에 직접 영향 받음.",
            "[A2A] 디버깅 어려움 — 여러 Agent 가 자율 결정해 추적 복잡.",
            "[A2A] 비용 ↑ — Agent 마다 LLM 호출 누적.",
            "[A2A] 표준 부족 — 아직 industry standard 정립 중.",
        ],
        "강추": [
            "보통 = MCP — 단일 LLM 으로 도구 호출, 시작 쉬움.",
            "멀티 도메인 자율 협력 = A2A — 큰 시스템.",
            "둘 다 혼합 가능 — MCP 로 도구 + A2A 로 Agent 협력 (우리 시스템 본보기).",
            "PoC = MCP 만, 프로덕션 = MCP + A2A.",
            "단순 작업 = MCP, 복잡 워크플로우 = A2A.",
        ],
        "우리시스템": {
            "MCP": "GitHub·Playwright·Figma·Slack·Notion MCP 서버 설치·사용",
            "A2A 스타일": "Claude (설계) → Codex (구현) → Gemini (검증) → Haiku (판정) 인수인계 (hook-08-ai-handoff)",
            "결론": "MCP + A2A 동시 사용 (다양한 작업에 적합한 방식 선택)",
        },
        "점검": ("Slack 으로 메시지 보내기 = MCP or A2A?", "MCP. 외부 도구 호출이라."),
    },
    # ---- 13. Claude 14 Levels ----
    {
        "title": "13. Claude 마스터 로드맵 14 단계",
        "image_eng": "클로드마스터-로드맵-14레벨-RubenHassid.jpg",
        "image_kor": "07-14-levels.png",
        "핵심": "Claude 를 0에서 마스터까지 14 단계. Lv5 까지만 가도 일상 작업 80% 자동화.",
        "강사": ["Claude Code 사용 능력 = 14 레벨로 나눌 수 있어요.", "Lv 1-5 (입문): 단일 명령·기본 hook.", "Lv 6-10 (중급): MCP·Skills·Subagents 활용.", "Lv 11-14 (마스터): plugin 개발·multi-agent chain·CI/CD 통합.", "우리 시스템 사용자 = Lv 9-10 추정. 본 강의 끝나면 Lv 11+ 가능."],
        "표": {
            "header": ["Lv", "단계", "핵심 결과"],
            "rows": [
                ["1-2", "가입·모델 선택", "Pro $20 결제, Opus 4.7"],
                ["3", "프롬프트 잘쓰기", "구체적·예시·새 대화"],
                ["4", "도구 연결", "Gmail·Drive·Gamma 등"],
                ["5", "Cowork", "PC 폴더 직접 읽고 씀"],
                ["6", "컨텍스트 폴더", "about-me 등 2000단어 이하"],
                ["7", "음성 입력", "WisprFlow"],
                ["8", "Obsidian", "마크다운 관리"],
                ["9", "Skills", "/linkedin·/negotiation"],
                ["10", "프로젝트", "세션 간 메모리"],
                ["11-12", "도구 추가·토큰 절약", ""],
                ["13", "팀 배포", "Shared Projects"],
                ["14", "본인 안목 ★", "AI 가 못 뺏는 영역"],
            ],
        },
        "흐름": [
            "Lv 1-3 (가입·모델·프롬프트) 첫 주 — 기본기.",
            "Lv 4-5 (도구 연결·Cowork) 둘째 주 — 일상 작업 80% 자동화 시작.",
            "Lv 6-10 (컨텍스트·Skills·프로젝트) 1-2개월 — 본격 사용.",
            "Lv 11-13 (도구 추가·토큰 절약·팀 배포) 3-6개월 — 시스템 구축.",
            "Lv 14 (본인 안목) 평생 — AI 가 못 뺏는 영역.",
        ],
        "강점": [
            "로드맵화 — 막막함 ↓, 단계별 성취감.",
            "각 Lv 가 독립적 — Lv5 만 가도 큰 효과.",
            "단계별 시간 투자 ↑ — 한 번에 다 안 해도 됨.",
            "Lv 9 (Skills) 부터는 본인만의 도구 — 차별화.",
            "Lv 14 (안목) 가 사람 영역 — AI 시대에도 사람 가치 ↑.",
        ],
        "약점": [
            "Lv 13-14 는 팀 환경 필요 — 혼자면 Lv 5-10 에서 만족.",
            "Lv 11+ 는 유료 도구 비용 ↑ — Excel·Design·Code Pro.",
            "Lv 6 (컨텍스트 폴더) 는 만들기 노동 — 작성 시간 큼.",
            "Lv 7 (음성) 은 한국어 환경에서 정확도 제한.",
            "Lv 마다 학습 곡선 — 한 번에 너무 많은 도구 X.",
        ],
        "강추": [
            "오늘 시작이면 Lv 1-3 (가입·모델·프롬프트) 만 — 1주.",
            "Lv 4-5 (도구 연결·Cowork) 추가 — 2주.",
            "Lv 9 (Skills) 가 게임 체인저 — 반복 작업 자동화.",
            "Lv 13 (팀 배포) 는 우리 install/setup 패턴 활용.",
            "Lv 14 는 항상 — AI 출력 중 어느 걸 출시할지 본인 안목.",
        ],
        "우리시스템": {
            "현재 사용자 위치": "Lv 9-10 추정 (Skills + 프로젝트 활용 중)",
            "Lv 11 추가도구": "Excel/Design/Code 다양 사용",
            "Lv 13 팀 배포": "orchestration_v1 의 install/setup 으로 팀 배포 가능 (template kit)",
            "Lv 14 본인 안목": "여전히 사람의 영역 — 깐깐 검수도 사람이",
        },
        "점검": ("Lv 14 (본인 안목) 가 가장 중요한 이유?", "AI 가 10개 만들면 어느 걸 선택할지는 사람의 가치판단. 대체 불가."),
    },
    # ---- 14. Decision Tree ----
    {
        "title": "14. Claude Code 결정트리 — Skills/Subagents/MCP/Hooks",
        "image_eng": "클로드코드-아키텍처-결정트리-판데이.jpg",
        "image_kor": None,
        "핵심": "4 가지 도구 (Skills/Subagents/MCP/Hooks) 중 무엇을 쓸지 한 그림에 결정.",
        "강사": ["4 도구 (Skills/Subagents/MCP/Hooks) 중 무엇을 쓸지 결정:", "지식 = Skills/CLAUDE.md. 외부 호출 = MCP. 강제 = Hooks. 격리 추론 = Subagent.", "**위험 차단 = 반드시 Hook** (AI 신뢰 X). rm -rf 차단은 PreToolUse hook 강제.", "우리 = 4/4 다 활용."],
        "표": {
            "header": ["도구", "트리거", "특징", "예"],
            "rows": [
                ["Skills", "모델이 결정 (description)", "지식·작업별", "/pr-review-checklist"],
                ["Subagents", "명시적 호출", "격리 추론·자체 모델", "/agent code-reviewer"],
                ["MCP", "모델이 결정", "외부 시스템 호출", "GitHub·Slack·Postgres"],
                ["Hooks", "이벤트 강제", "결정론·AI 아님", "PreToolUse·Stop"],
            ],
        },
        "흐름": [
            "Q1: 지식? 행동? → 지식 = Skills/CLAUDE.md.",
            "Q2 (행동): 내부 추론? 외부 호출? → 외부 = MCP or Hooks.",
            "Q3 (외부): 모델이 결정? 강제? → 모델결정 = MCP / 강제 = Hooks.",
        ],
        "강점": [
            "역할 분리로 디버깅 쉬움 — 어디서 결정됐는지 명확.",
            "Hooks 가 AI 아닌 결정론적 강제 — 가장 신뢰.",
            "Subagents 의 컨텍스트 격리로 메인 보호.",
            "MCP 표준으로 새 외부 시스템 추가 쉬움.",
            "Skills 자동 호출 — 사용자 명시 X.",
        ],
        "약점": [
            "Subagents 격리 = 부모 메모리 못 봄. 정보 전달 필요.",
            "Subagents 호출 비용 ↑ — 자체 모델 인스턴스.",
            "Hooks 디버깅 어려움 — 자동 발동이라 추적 도구 필요.",
            "MCP 의 외부 시스템 의존 — 네트워크·인증 fail 가능.",
            "Skills description 매칭 실패 시 호출 X — false negative.",
        ],
        "강추": [
            "위험 차단 = 반드시 Hook (모델 신뢰 X).",
            "큰 코드 탐색 = Explorer subagent.",
            "외부 SaaS = MCP.",
            "재사용 워크플로우 = Skill.",
        ],
        "우리시스템": {
            "Skills": ".claude/skills/ 77개 (auto-compact·llm-as-judge·route_dispatch 등)",
            "Subagents": ".claude/agents/ 11개 (team-lead·reviewer·judge·explorer·test-runner)",
            "MCP": "GitHub·Playwright·Figma·Slack·Notion·Mermaid 등 다수",
            "Hooks": "settings.json 24개 등록 (block_dangerous_bash·protect-critical-files 등)",
        },
        "점검": ("git rm -rf 자동 차단은 어느 도구?", "Hook (PreToolUse Bash). 모델 결정 X."),
    },
    # ---- 15. Complete Guide ----
    {
        "title": "15. Claude Code 완전 가이드 — 6 단계 사이클",
        "image_eng": "클로드코드-완전가이드-판데이.jpg",
        "image_kor": None,
        "핵심": "Install → Configure → Prompt → Review → Iterate → Ship. 6 단계 사이클이 끝까지 반복.",
        "강사": ["6 단계 사이클: Install → Configure → Prompt → Review → Iterate → Ship.", "각 단계 산출물 명확. 작은 PR 단위로 사이클 돌리세요.", "Review·Iterate 가 핵심 — 한 번에 완벽 X.", "우리 = 6/6 다 작동. install.bat (Zero-touch) + eval_quality + watchdog + auto-review CI."],
        "표": {
            "header": ["단계", "할 일", "산출물"],
            "rows": [
                ["1 Install", "Claude Code CLI 설치", "claude 명령"],
                ["2 Configure", "CLAUDE.md·hooks·skills", ".claude/ 폴더"],
                ["3 Prompt", "구체적 요청 작성", "task-instruction"],
                ["4 Review", "결과 검토·수정 요청", "리뷰 코멘트"],
                ["5 Iterate", "반복 개선", "v2, v3..."],
                ["6 Ship", "출시·배포", "merged PR"],
            ],
        },
        "흐름": [
            "1 Install → 2 Configure → 3 Prompt → 4 Review → 5 Iterate → 6 Ship → (필요시 3 으로 루프).",
            "각 단계마다 산출물 명확 — claude 명령, .claude/ 폴더, task-instruction, 리뷰 코멘트, v2, merged PR.",
            "4 Review 와 5 Iterate 가 핵심 — 한 번에 완벽 X, 점진 개선.",
        ],
        "강점": [
            "사이클 명확 — 어느 단계서 막혀도 다음 행동 알 수 있음.",
            "각 단계 산출물 정의 — 진행 상황 가시화.",
            "Review·Iterate 가 품질 보장 — 한 번에 완벽 안 해도 됨.",
            "Install/Configure 한 번이면 끝 — 매번 반복 X.",
            "Ship 후에도 3 으로 루프 가능 — 기능 추가·버그 fix.",
        ],
        "약점": [
            "사이클마다 사람 개입 필요 — 자동화 한계 (특히 Review).",
            "Configure 가 초기 노동 ↑ — CLAUDE.md, hooks, skills 작성.",
            "Iterate 가 무한 루프 위험 — 완벽주의 함정.",
            "Ship 기준이 모호 — 어느 시점에 출시할지 본인 판단.",
            "작은 변경에도 1-6 다 도는 게 과함 — 핫픽스 fast path 별도 필요.",
        ],
        "강추": [
            "처음엔 1-3 만 익히세요 — 4-6 은 자연스럽게 익숙.",
            "작은 PR 단위로 사이클 돌리세요 — 큰 변경은 위험.",
            "Review 단계에 eval_quality + score-task 활용 — 자동 채점.",
            "Iterate 는 최대 3회 — 더 돌면 처음부터 다시 (Prompt).",
            "Ship 기준 = 핵심 테스트 통과 + 리뷰 OK + 비용 임계 이내.",
        ],
        "우리시스템": {
            "1 Install": "install.bat / setup/setup.bat 1회 실행 (zero-touch 보강 중)",
            "2 Configure": "이미 설정됨 (CLAUDE.md 169줄, 24 hooks, 77 skills)",
            "3 Prompt": "사용자 채팅 또는 task-instruction.md",
            "4 Review": "eval_quality + score-task (Haiku 자동 채점)",
            "5 Iterate": "watchdog 가 죽은 워커 재시작 = 자동 iterate",
            "6 Ship": "git commit + (선택) gh pr create",
        },
        "점검": ("이 사이클에서 hook 이 자동 발동되는 단계?", "2-6 전체 (PreToolUse·PostToolUse·Stop 등)."),
    },
    # ---- 16. Architecture Reference ----
    {
        "title": "16. Claude Code 아키텍처 레퍼런스 — 60초 셋업",
        "image_eng": "클로드코드-아키텍처-레퍼런스-판데이.jpg",
        "image_kor": None,
        "핵심": "한 장에 5 레이어 + 60초 셋업 + 키 단축키. 두고두고 보는 치트시트.",
        "강사": ["한 장에 5 레이어 + 60초 셋업 + 키 단축키 정리한 치트시트.", "60초 셋업: npm install -g claude-code → claude → /init → 끝.", "/compact (컨텍스트 압축), /clear (초기화) 외우면 비용 ↓.", "주 1회 봐주세요 — 새 단축키 발견 = 시간 절약."],
        "표": {
            "header": ["Layer", "이름", "위치"],
            "rows": [
                ["Foundation", "Runtime", "claude 명령"],
                ["1", "Memory System", "CLAUDE.md (3 scopes)"],
                ["2", "Skills Engine", ".claude/skills/"],
                ["3", "MCP Connections", "200+ 도구"],
                ["4", "Commands", "/init·/clear·/compact"],
                ["5", "Orchestration", "Hooks·체크포인트"],
            ],
        },
        "흐름": [
            "60초 셋업: npm install -g claude-code → claude 명령 → /init → 끝.",
            "Layer 1 (Memory) 부터 차례로 — CLAUDE.md 정의 → Skills 추가 → MCP 연결 → Commands 활용 → Hooks 설정.",
            "각 Layer 가 다음 Layer 의 기반 — 1 없이 2 만 만들면 컨텍스트 부재.",
        ],
        "강점": [
            "치트시트 화 가능 — 모니터 옆에 출력 붙여둘 만함.",
            "5 레이어 전체 한 장에 정리 — 큰 그림 명확.",
            "60초 셋업 — 진입 장벽 ↓.",
            "키 단축키 (Esc·Shift+Esc 등) 포함 — 시간 절약.",
            "60-second setup 으로 빠른 PoC 가능.",
        ],
        "약점": [
            "정보 밀도 ↑ — 처음엔 어디부터 봐야 할지 모름.",
            "Foundation Runtime 은 OS 별 차이 — Windows·Mac·Linux 별도 확인.",
            "Layer 4-5 (Commands·Orchestration) 는 사용해봐야 이해.",
            "치트시트 1장에 모두 압축이라 깊이 부족.",
            "버전마다 단축키·명령 바뀜 — 최신 버전 확인 필요.",
        ],
        "강추": [
            "주 1회 봐주세요 — 새 키 단축키 발견 = 시간 절약.",
            "출력해서 모니터 옆에 — 자주 보면 외워짐.",
            "Layer 별로 점검 — 1·2 가 약하면 3-5 효율 ↓.",
            "/compact, /clear, /reset context 같은 명령 외우면 비용 ↓.",
        ],
        "우리시스템": {
            "Foundation": "claude --dangerously-skip-permissions 자동",
            "Layer 1-5": "위 챕터 6 의 5 레이어와 동일 매핑",
        },
        "점검": ("\"내 컨텍스트 너무 길어\" 할 때 명령?", "/compact. 자동 압축."),
    },
    # ---- 17. Project Structure ----
    {
        "title": "17. Claude Code 프로젝트 구조",
        "image_eng": "클로드코드-프로젝트구조-판데이.jpg",
        "image_kor": None,
        "핵심": "어디에 무엇을 두는가 표준. 처음부터 이 구조면 90점.",
        "강사": ["어디에 무엇을 두는가 표준 — 90점 시작.", "CLAUDE.md (팀 규칙) · settings.json (권한+hook) · .claude/ (commands/skills/agents/hooks) · plugins/ (원본 SoT).", "이 구조만 따르면 90점. 나머지 10점 = 팀 합의로 보완."],
        "표": {
            "header": ["폴더", "역할", "예"],
            "rows": [
                ["루트 CLAUDE.md", "팀 공유 규칙", "git commit"],
                [".claude/commands/", "슬래시 명령", "/review·/test"],
                [".claude/skills/", "자동 워크플로우", "SKILL.md"],
                [".claude/agents/", "subagent", "code-reviewer.md"],
                [".claude/hooks/", "가드레일", "PreToolUse.sh"],
                [".mcp.json", "MCP 서버 설정", "GitHub·Postgres"],
            ],
        },
        "흐름": [
            "새 프로젝트 = /init → 위 구조 자동 생성 → 필요한 파일만 채움.",
            "CLAUDE.md 부터 정의 → skills/ 추가 → agents/ 위임 → hooks/ 가드레일 → .mcp.json 외부 도구.",
            "plugins/ 는 sync 원본 — .claude/ 는 sync 결과물 (자동 생성).",
        ],
        "강점": [
            "표준 폴더 — 다른 사람이 봐도 즉시 이해.",
            "/init 한 번이면 끝 — 자동 생성.",
            "각 폴더 책임 분명 — 어디에 무엇 둬야 할지 헷갈림 X.",
            "Git 친화 — CLAUDE.md commit, settings.local gitignore.",
            "팀 표준화 쉬움 — 모두 같은 구조 따름.",
        ],
        "약점": [
            "폴더 많으면 처음 부담 — 필요한 것만.",
            "plugins/ 와 .claude/ 분리 → 어디 편집할지 헷갈림 (plugins/ 가 SoT).",
            "settings.local 잊고 commit 위험 — 시크릿 누출.",
            "agents/ 와 skills/ 차이 헷갈림 — subagent 는 명시 호출, skill 은 자동.",
        ],
        "강추": [
            "팀 작업 = 무조건 이 구조 — 표준 준수.",
            "혼자면 CLAUDE.md + skills/ 만으로 시작 — 점진 확장.",
            "/init 먼저 → 큐레이션 — Anthropic 권장 워크플로우.",
            "sync 충돌 방지 = plugins/ 에서만 편집.",
        ],
        "우리시스템": {
            ".claude/commands/": "152개 (sync 결과물)",
            ".claude/skills/": "77개",
            ".claude/agents/": "11개 (최근 +4 sync 보강)",
            ".claude/hooks/": "21 .sh/.py + settings.json 등록",
            "plugins/": "25개 (원본 SoT)",
        },
        "점검": ("CLAUDE.md 와 CLAUDE.local.md 차이?", "전자 = 팀 공유 (commit), 후자 = 개인용 (gitignore)."),
    },
    # ---- 18. DK Folder Structure ----
    {
        "title": "18. .claude 폴더 전체 구조 (DK 메소드)",
        "image_eng": "클로드폴더-전체구조-DK메소드.jpg",
        "image_kor": None,
        "핵심": "한글로 가장 친절하게 정리된 .claude/ 트리.",
        "강사": [".claude 폴더 = AI 팀원의 사무실.", "commit (팀 공유): CLAUDE.md · settings.json · commands/ · rules/ · skills/ · agents/", "gitignore (개인): CLAUDE.local.md · settings.local · .env", "시크릿은 절대 commit X — 우리 secret-scan hook 자동 차단."],
        "표": {
            "header": ["파일/폴더", "용도", "git"],
            "rows": [
                ["CLAUDE.md", "팀 공유 기억", "commit"],
                ["CLAUDE.local.md", "나만의 기억", "gitignore"],
                [".claude/settings.json", "권한 설정 (공유)", "commit"],
                [".claude/settings.local", "개인 권한", "gitignore"],
                [".claude/commands/", "나만의 /명령어", "commit"],
                [".claude/rules/", "항상 적용 규칙", "commit"],
                [".claude/skills/", "자동 워크플로우", "commit"],
                [".claude/agents/", "서브에이전트", "commit"],
            ],
        },
        "흐름": [
            "새 프로젝트 = /init → 위 파일들 만들어짐 → 팀과 합의로 채움.",
            "팀 공유 (CLAUDE.md, settings.json, commands/, rules/, skills/, agents/) 와 개인 (CLAUDE.local.md, settings.local, .env) 분리.",
            "settings.json 에 hooks 등록 → PreToolUse/PostToolUse 자동 발동.",
        ],
        "강점": [
            "공유 (commit) 와 개인 (gitignore) 명확 분리 — 시크릿 안전.",
            "한글 친화 — DK 메소드가 한글 라벨 풍부.",
            "rules/ 폴더로 항상 적용 규칙 별도 관리.",
            "팀 onboarding 쉬움 — 신입이 폴더 봐도 즉시 이해.",
            "Git 정책 명확 — commit/gitignore 표시.",
        ],
        "약점": [
            ".local 파일 잊고 commit 하면 시크릿 누출 위험 — pre-commit hook 필수.",
            "CLAUDE.md vs CLAUDE.local.md 혼동 — 어느 것에 뭐 쓸지 결정 필요.",
            "rules/ 와 skills/ 차이 헷갈림 — rules 는 규칙, skills 는 워크플로우.",
            "Agent 가 많으면 description 매칭 충돌.",
        ],
        "강추": [
            "팀 작업 = CLAUDE.md + .local.md 둘 다 활용 — 분담.",
            "시크릿은 무조건 .local 또는 .env — 절대 commit X.",
            "rules/ 에 code-style.md 같은 강제 룰 — Hook 으로 자동 검증.",
            "agents/ 에 code-reviewer, test-runner 같은 핵심 subagent 만.",
        ],
        "우리시스템": {
            "CLAUDE.md": "✓ 169줄, commit 됨",
            "CLAUDE.local.md": "사용 X (전체 가이드에 다 박음)",
            "settings.json": "24 hooks 등록",
            ".local 파일": ".env (gitignore) 로 시크릿 분리",
        },
        "점검": ("팀과 공유 안 할 메모는 어디에?", "CLAUDE.local.md (gitignore)."),
    },
    # ---- 19. CLAUDE.md design guide ----
    {
        "title": "19. CLAUDE.md 설계 가이드 — 진짜로 작동하는 메모리",
        "image_eng": "클로드MD-설계가이드-판데이.jpg",
        "image_kor": None,
        "핵심": "CLAUDE.md 는 사람용 README 가 아니라 AI 팀원 온보딩 문서.",
        "강사": ["CLAUDE.md 는 사람용 README 가 아니라 **AI 팀원 온보딩 문서**.", "3 Scope: Global (~/.claude/) · Project (./CLAUDE.md) · Folder (./src/CLAUDE.md). 가까운 게 이김.", "5 Rules: /init 먼저 · 500줄 이하 · Hooks 사용 · 월간 갱신 · 참조 중심.", "우리 = 169줄 (한계의 33%, 충분히 짧음) + 5중박기 + 13 금지."],
        "표": {
            "header": ["Scope", "위치", "용도"],
            "rows": [
                ["Global", "~/.claude/CLAUDE.md", "모든 프로젝트 공통"],
                ["Project", "./CLAUDE.md", "이 프로젝트 규칙"],
                ["Folder", "./src/CLAUDE.md", "모듈별 오버라이드"],
            ],
        },
        "흐름": [
            "WHAT (컨텍스트) → 목적·기술스택·구조",
            "WHY (원칙) → 아키텍처 결정·스타일·안티패턴",
            "HOW (워크플로우) → build·test·lint·commit",
        ],
        "강점": [
            "3 스코프로 충돌 해결 — Folder > Project > Global (가까운 게 이김).",
            "WHAT/WHY/HOW 프레임으로 빠진 정보 없이 정리.",
            "5 Rules 가 실용적 — /init 먼저, 500줄 이하, Hooks, 월간 갱신, 참조 중심.",
            "Hooks 와 결합으로 70% (메모리) → 100% (Hooks) 강제력.",
            "AI 팀원 온보딩 표준 — 사람용 README 아닌 AI 용 문서.",
        ],
        "약점": [
            "500줄 넘으면 무시됨 — 길어지지 않게 참조 중심.",
            "Folder scope 는 거의 안 쓰임 — 복잡도 ↑.",
            "Specific 하게 쓰기 어려움 — 'clean code' 같은 모호한 표현 함정.",
            "월간 갱신 노동 — 살아있는 문서 유지 부담.",
            "GLOBAL 의 시크릿 수정 시 모든 프로젝트 영향 — 위험.",
        ],
        "강추": [
            "1) /init 먼저 — Claude 가 잡게 둔 후 큐레이션",
            "2) 500줄 이하",
            "3) Hooks 사용 (메모리는 70%, hooks 는 100% 강제)",
            "4) 월간 업데이트",
            "5) 참조 중심, 중복 X",
        ],
        "우리시스템": {
            "Global": "~/.claude/CLAUDE.md (방금 install 자동 배포)",
            "Project": "./CLAUDE.md 169줄 (5중 박기 + 13개 금지)",
            "Folder": "사용 X",
            "강화 위치": "rules/teaching-doc.md·failure-mode.md·best-practices.md 로 분산",
        },
        "점검": ("CLAUDE.md 가 500줄 넘으면?", "무시됨. 참조 파일로 분산."),
    },
    # ---- 20. Prompt Frameworks ----
    {
        "title": "20. 8가지 프롬프트 프레임워크 — 결과를 두 배로",
        "image_eng": "클로드프롬프트-8가지프레임워크-natanmohart.jpg",
        "image_kor": None,
        "핵심": "프롬프트는 '대충 쓰기' 가 아니라 '템플릿'. 8 중 한 개만 외워도 80% 커버.",
        "강사": ["프롬프트는 '대충 쓰기' 아니라 '템플릿'. 8 중 1개 (CLARITY) 만 외워도 80%.", "CLARITY = Context+Look+Ask+Rules+Input+Target+Yardstick.", "복잡 작업 = TRUST. 데이터 분석 = RIPPLE.", "우리 task-instruction.md = CLARITY 구조와 일치. 다 활용."],
        "표": {
            "header": ["#", "프레임워크", "쓰는 곳"],
            "rows": [
                ["1", "CLARITY", "처음 시작·만능"],
                ["2", "SOCRATES", "단계별 계획"],
                ["3", "ANTICIPATE", "상품 기획"],
                ["4", "PARTNER", "콘텐츠 전략"],
                ["5", "TRUST", "깊이 있는 분석"],
                ["6", "RIPPLE", "데이터 분석"],
                ["7", "CATCH", "마케팅 카피"],
                ["8", "MAGIC", "랜딩페이지"],
            ],
        },
        "흐름": [
            "요청 유형 분류 (콘텐츠·계획·분석·마케팅) → 적합 프레임워크 선택.",
            "프레임워크 슬롯 채워서 프롬프트 작성 — 빠진 정보 없게.",
            "한 번 작성하면 템플릿으로 저장 → 재사용.",
            "프레임워크는 가이드라인 — 상황 따라 조정 (모든 슬롯 안 채워도 OK).",
        ],
        "강점": [
            "프롬프트 품질 일관 ↑ — 매번 다른 결과 X.",
            "복붙해서 재사용 가능 — 시간 절약.",
            "구조화되어 LLM 이 이해하기 쉬움.",
            "팀 표준화 가능 — 같은 작업에 같은 프레임워크.",
            "빠진 정보 자동 인식 — 슬롯이 체크리스트 역할.",
        ],
        "약점": [
            "프레임워크 너무 길면 LLM 컨텍스트 잡아먹음 — 토큰 비용 ↑.",
            "8개 다 외우면 결정 마비 — 1개만으로 충분.",
            "프레임워크에 갇히면 창의성 ↓ — 가끔 자유 형식도 필요.",
            "상황별 적합도 다름 — 마케팅에 SOCRATES 쓰면 과함.",
            "한국어 매칭 안 됨 — 영어 약자 (CLARITY 등) 가 한국어 단어 안 됨.",
        ],
        "강추": [
            "CLARITY 한 개만 외우세요 — 80% 커버.",
            "단계별 계획 = SOCRATES.",
            "콘텐츠·마케팅 = CATCH or MAGIC.",
            "데이터 분석 = RIPPLE.",
            "복잡 작업 = TRUST.",
        ],
        "우리시스템": {
            "CLARITY": "task-instruction.md 의 구조와 비슷 (Context + Rules + Target + Constraints)",
            "TRUST": "code review·아키텍처 결정 시 적합",
            "5중 박기 원칙": "결과적으로 SOCRATES·ANTICIPATE 와 비슷한 효과",
        },
        "점검": ("프레임워크 8개 다 외워야 하나?", "X. CLARITY 한 개만으로 80% 커버."),
    },
]


# ============================================================
# 문서 생성
# ============================================================
doc = Document()

# 페이지: landscape (가로) — 이미지 글씨 최대 가독성
from docx.enum.section import WD_ORIENT
from docx.shared import Mm
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    # A4 landscape (297×210mm = 11.69×8.27 inch) — PNG 비율 0.69 와 inside 비율 0.70 일치
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Cm(0.3)
    section.bottom_margin = Cm(0.3)
    section.left_margin = Cm(0.3)
    section.right_margin = Cm(0.3)

# default view zoom 100% + 인쇄 레이아웃 강제 (Word 가 페이지 가득 표시)
# w:val="none" 강제 — "bestFit"이면 percent 무시되고 창에 맞춤 (작아 보임)
_settings = doc.settings.element
_zoom = _settings.find(qn("w:zoom"))
if _zoom is None:
    _zoom = OxmlElement("w:zoom")
    _settings.insert(0, _zoom)
_zoom.set(qn("w:val"), "none")
_zoom.set(qn("w:percent"), "100")
_view = _settings.find(qn("w:view"))
if _view is None:
    _view = OxmlElement("w:view")
    _settings.insert(0, _view)
_view.set(qn("w:val"), "print")


# ---- 표지 ----
P(doc, "", size=8, after=80)
P(doc, "AI 에이전트 + Claude Code", size=30, bold=True, align="center",
  color=(32, 56, 100), after=8)
P(doc, "초보자 가이드 — 24장 그림으로 끝내는 2026 AI 입문",
  size=14, align="center", color=(89, 89, 89), after=40)
P(doc, "강사: AI 에이전트 강사 (Claude)", size=12, align="center",
  color=(120, 120, 120), after=4)
P(doc, "교재: docs/screens/arch (영어 원본) + docs/screens/arch-kor (한글 다이어그램)",
  size=10, align="center", color=(150, 150, 150), after=4)
P(doc, "버전: 2026-05-11 v2 (8섹션 표준 + 한글 다이어그램)",
  size=10, align="center", color=(150, 150, 150), after=60)

P(doc, "▷ 이 책의 약속 (8섹션 표준)", size=13, bold=True, color=(192, 0, 0), after=6)
B(doc, "각 챕터에 표·흐름·강점·약점·강추·우리 시스템 매핑·점검 모두 포함")
B(doc, "외국어 그림은 한글 다이어그램과 같이 (matplotlib 자동 생성)")
B(doc, "5살 청자 톤 — 어려운 단어 즉시 풀이, 비유 풍부")
B(doc, "막히면 부록 B '한 줄 정리' 부터 30분 안에 전체 파악")
PB(doc)


# ---- 목차 ----
H(doc, "목차", level=1)
B(doc, "0. 들어가며")
for i, ch in enumerate(CHAPTERS, 1):
    B(doc, ch["title"])
B(doc, "부록 A. 한 줄 정리")
B(doc, "부록 B. 고퀄리티 다이어그램 도구 (Canva·Figma·Mermaid)")


# ---- 0. 들어가며 (IMG paragraph 자체에 page_break_before — 빈 PB 페이지 방지) ----
_intro_para = doc.add_paragraph()
_intro_para.paragraph_format.page_break_before = True
_intro_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
_intro_para.paragraph_format.space_after = Pt(0)
from PIL import Image as _PILImage
_img_path = ARCH_KOR / "00-ai-evolution.png"
with _PILImage.open(str(_img_path)) as _pim:
    _iw, _ih = _pim.size
_w_in, _h_in = 11.46, 11.46 * (_ih / _iw)
if _h_in > 8.0:
    _h_in = 8.0
    _w_in = 8.0 * (_iw / _ih)
_intro_para.add_run().add_picture(str(_img_path), width=Inches(_w_in), height=Inches(_h_in))


# ---- 19 챕터 자동 렌더링 ----
for idx, ch in enumerate(CHAPTERS):
    render_chapter(doc, ch, idx=idx)


# ---- 부록 A — 한 줄 정리 (이전 부록 B) ----
H(doc, "부록 A. 한 줄 정리 — 5분 안에 전체 복습", level=1)
cheats = [
    ("Generative vs Agentic vs Agent", "글 → 단계 → 손까지. 셋이 같은 AI 아님."),
    ("AI 8가지 모델", "GPT 만 있지 않음. 작업에 맞는 뇌 골라쓰기."),
    ("에이전트 5 핵심", "Guardrail / Orch / MCP / Memory / Observability."),
    ("9 Silent Killers", "#2 맥락 부패와 #9 거짓말 두 개만 잡아도 신뢰도 두 배."),
    ("AI 스택 5층", "Infra → Data → LLM → Orch → Interface."),
    ("개발킷 5레이어", "CLAUDE.md + Skills + Hooks + Subagents + Plugins."),
    ("$0 AI 2026", "Streamlit + Ollama + ChromaDB 30 분 안에 띄움."),
    ("AI 빌더 30 도구", "초보 = 카테고리 1 (Claude) + 2 (Claude Code) 만."),
    ("RAG", "사내 문서 묻고 답할 때. 처음엔 Naive."),
    ("RAG 8가지", "Naive → Corrective → Graph → Agentic 순으로 발전."),
    ("API 프로토콜", "REST + Webhooks 만 알면 80% 끝."),
    ("MCP vs A2A", "MCP = 도구 호출. A2A = 에이전트 협력."),
    ("Claude 14 레벨", "Lv5 (Cowork) 까지면 일상 자동화 80%."),
    ("결정트리", "지식 = Skill/CLAUDE.md. 행동 = Subagent/MCP/Hook."),
    ("완전 가이드", "Install → Configure → Prompt → Review → Iterate → Ship."),
    ("60초 셋업", "/init 만 치면 폴더 잡힘."),
    ("프로젝트 구조", "CLAUDE.md + .claude/{commands,skills,agents,hooks}."),
    (".claude DK", "CLAUDE.md(공유) + CLAUDE.local.md(개인) 분리."),
    ("CLAUDE.md", "500줄 이하·Hooks 사용·월간 갱신·참조 중심."),
    ("프롬프트", "CLARITY 한 개만 외우면 80% 커버."),
]
for k, v in cheats:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    _set_kor(p.add_run("◆ " + k + " — "), size=11, bold=True, color=(192, 0, 0))
    _set_kor(p.add_run(v), size=11)
PB(doc)


# ---- 부록 B — 고퀄리티 도구 (이전 부록 C) ----
H(doc, "부록 B. 고퀄리티 다이어그램 도구 (가입 + MCP 등록)", level=1)
callout(doc, "📚 핵심 한 줄",
        "이 책 한글 다이어그램은 matplotlib 무료. 더 예쁘게 = 아래 도구 가입 + MCP 등록.")
quality_tools = [
    ("Mermaid Chart (mermaid.live)", "무료 / 유료 $14/월",
     "코드로 다이어그램. 한글 OK. 빠른 작업.", "claude.ai → 커넥터 → Mermaid 활성화"),
    ("Excalidraw", "무료 / 유료 $7/월",
     "손그림 톤. 발표용 좋음.", "https://excalidraw.com (API 유료)"),
    ("Figma", "무료 / 유료 $12/월",
     "전문 디자인 표준.", "/install-mcp 또는 Figma MCP + PAT"),
    ("Canva", "무료 / Pro $13/월",
     "인포그래픽·소셜·프레젠테이션. 한글 폰트 풍부.", "claude.ai → 커넥터 → Canva (1-click)"),
    ("Gamma", "무료 (제한) / Plus $10/월",
     "AI 프레젠테이션 자동 생성.", "claude.ai → 커넥터 → Gamma 활성화"),
]
for name, price, desc, how in quality_tools:
    P(doc, "◆ " + name, size=12, bold=True, color=(32, 56, 100), after=2)
    callout(doc, "가격:", price, label_color=(192, 0, 0))
    callout(doc, "용도:", desc, label_color=(89, 89, 89))
    callout(doc, "연결법:", how, label_color=(11, 83, 148))

H(doc, "💡 권장 조합", level=3)
B(doc, "초보: matplotlib (이 책) + Mermaid 무료.")
B(doc, "중급: Canva 무료 + Gamma 무료 (AI 프레젠테이션).")
B(doc, "전문: Figma Pro + Canva Pro + Mermaid Pro.")
H(doc, "🛠 우리 시스템 활용", level=3)
B(doc, "/install-mcp — Canva (OAuth) + Figma (PAT) + Mermaid 한 번에")
B(doc, "/plug_design — 디자인 MCP 패키지 (Canva·Figma·Gamma·PPT·Slides·Mermaid)")
B(doc, "/arch-auto, /arch-mindmap, /arch-layered, /arch-cheatsheet — 자동 다이어그램")


# ---- 마지막 ----
HR(doc)
P(doc, "끝까지 읽어주셔서 감사합니다.", size=12, align="center", bold=True, after=4)
P(doc, "막히면 Claude 에게 '5살한테 설명하듯 알려줘' 라고 부탁하세요.",
  size=10, align="center", color=(120, 120, 120))


# ---- 저장: .bak 백업 + 폴링 대기 (rule: 멈추지 마라, 자동 우회) ----
import time as _time
OUT.parent.mkdir(parents=True, exist_ok=True)


def _wait_unlock(path, max_sec=60, interval=2):
    """잠금 해제 대기 — 사용자가 Word 닫을 때까지 폴링."""
    elapsed = 0
    while elapsed < max_sec:
        try:
            test = path.with_suffix(path.suffix + ".lock-test")
            path.rename(test)
            test.rename(path)
            return True
        except (PermissionError, OSError):
            if elapsed == 0:
                print(f"[WAIT] {path.name} 잠김 — 자동 폴링 ({max_sec}초). 다른 작업 진행 가능.")
            _time.sleep(interval)
            elapsed += interval
    return False


if OUT.exists():
    bak = OUT.with_suffix(OUT.suffix + ".bak")
    try:
        bak.unlink(missing_ok=True)
        OUT.rename(bak)
    except PermissionError:
        if _wait_unlock(OUT, max_sec=60):
            bak.unlink(missing_ok=True)
            OUT.rename(bak)
        else:
            print(f"[FAIL] 60초 지나도 {OUT.name} 잠김. Word 직접 닫고 재실행.")
            raise SystemExit(1)

try:
    doc.save(str(OUT))
except PermissionError:
    if _wait_unlock(OUT, max_sec=60):
        doc.save(str(OUT))
    else:
        print(f"[FAIL] 저장 실패. Word 닫고 재실행.")
        raise SystemExit(1)

print(f"[OK] {OUT.name}  {OUT.stat().st_size:,} bytes")
print(f"     {len(CHAPTERS)} 챕터 × 8 섹션 표준")
if (OUT.with_suffix(OUT.suffix + ".bak")).exists():
    print(f"     백업: {OUT.name}.bak (이전 버전)")
